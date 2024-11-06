# 根据产品简述和产品需求分解excel,构建agent生成需求文档

import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import bs4
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import docx
from typing import Any
import logging


# Add after imports, before class definition
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = "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"
os.environ["OPENAI_API_KEY"] = "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"


class DocumentAgent:
    def __init__(self, overview_path, excel_path, docx_name, title):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.excel_path = excel_path
        self.docx_name = docx_name
        self.title = title
        
        # 定义产品综述的Prompt模板
        self.summary_prompt = PromptTemplate(
            input_variables=["product_description", "idx"],
            template=(
                "以下是一个产品的简述: {product_description}\n"
                "请基于这个简述生成产品需求文档的开头综述。"
                "概述该产品的背景与目标, 产品愿景, 适用范围, 不少于500字。"
                "这是整个word文档的第 {idx} 段。请在开头标明序号"
                "不需要在开头加产品需求文档, 标题已经有了"
            )
        )

        # 定义结论的Prompt模板
        self.conclusion_prompt = PromptTemplate(
            input_variables=["product_description", "idx"],
            template=(
                "基于以下的产品描述，请生成需求文档的结论部分，总结主要特点和应用价值, 不少于500字。"
                "这是产品描述: {product_description}"
                "这是整个word文档的第 {idx} 段。请在开头标明序号。"
            )
        )

    def generate_product_summary(self, product_description, idx=1):
        # 使用LLMChain生成产品综述
        self.summary_chain = LLMChain(
            llm=self.llm,
            prompt=self.summary_prompt
        )
        product_summary = self.summary_chain.run(product_description=product_description, idx=1)
        logger.info("Product Summary:", product_summary)
        return f"第 {idx} 段: {product_summary}"

    def generate_conclusion(self, product_description, idx):
        # 使用LLMChain生成结论
        self.conclusion_chain = LLMChain(
            llm=self.llm,
            prompt=self.conclusion_prompt
        )
        conclusion = self.conclusion_chain.run(product_description=product_description, idx=idx)
        logger.info("Document Conclusion:", conclusion)
        return conclusion

    def generate_module_content(self, module, module_background, idx): 
        # 模块内容生成的Prompt模板
        question_prompt = PromptTemplate(
            input_variables=["module", "module_background", "idx"],
            template=(
                "{module_background}。\n"
                "基于产品综述和功能需求信息，请为模块{module}生成产品需求文档的段落, 不少于500字。"
                "段落应按顺序编号, 包含详细需求和设计思路, 以便最终输出到Word文档。"
                "这是整个word文档的第 {idx} 段, 请在开头标明序号。注意生成的是整个文档的一个段落"
            )
        )
        
        module_chain = LLMChain(
            llm=self.llm,
            prompt=question_prompt
        )
        
        content = module_chain.run(module=module, module_background=module_background, idx=idx)
        logger.info(f"生成的模块{module}内容: {content}")
        return content

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()

        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        # 生成并添加产品总结
        product_summary = self.generate_product_summary(product_description, idx=1)
        document.add_paragraph(product_summary)

        # 从Excel中读取模块信息并存储到Chroma DB
        sheets = pd.read_excel(self.excel_path, sheet_name=None)
        df = sheets[next(iter(sheets))]
        self.modules = df['模块'].unique().tolist()
        grouped = df.groupby('模块')['功能'].apply(lambda x: ', '.join(x)).reset_index()
        for _, row in grouped.iterrows():
            module = row['模块']
            functions = row['功能']
            doc_content = f"这是一个关于{module}模块的文档，其中功能包括: {functions}。"
            self.chroma_db.add_texts([doc_content])

        # 检索并生成每个模块的内容，保存到文档中
        idx = 2
        for module in self.modules:
            query = f"关于{module}"
            docs = self.chroma_db.similarity_search(query)
            filtered_docs = [doc for doc in docs if module in doc.page_content]

            if len(filtered_docs) == 0:
                logger.warning(f"模块{module}没有检索到相关内容")
                continue
            if len(filtered_docs) > 1:
                logger.warning(f"模块{module}检索到多条数据，可能存在重复")
                continue
            
            doc = filtered_docs[0]
            module_background = f"产品综述: {product_description}。 功能需求信息: {doc}"
            module_content = self.generate_module_content(module, module_background, idx)

            
            document.add_paragraph(module_content)
            idx += 1

        # 生成并添加结论
        conclusion = self.generate_conclusion(product_description, idx)
        
        document.add_paragraph(conclusion)

        # 保存文档
        document.save(self.docx_name)
        logger.info(f"需求文档已保存为 {self.docx_name}")

if __name__ == "__main__":
    agent = DocumentAgent(
        "./大模型智能营销系统功能说明.txt",
        "./大模型智能营销系统.xlsx",
        "大模型智能营销系统功能说明.docx",
        "大模型智能营销系统需求文档"
    )
    agent.process()
