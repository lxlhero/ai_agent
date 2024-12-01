import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
import docx
from typing import Any
import logging
from docx.shared import Pt
import traceback
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment



# 数据集技术规范书生成器
# 根据数据集简述生成技术规范书



# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, dataset_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        # 这里是整个产品的名称
        self.dataset_name = dataset_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的专业数据集描述，撰写技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            撰写的内容必须遵循以下要求：
            范围明确：本次目的是为一个新构建的数据集提供技术规范书，所以不要针对一类数据集泛泛而谈，而是聚焦于描述一个特定数据集的技术规范, 相关的技术标准可以参考该领域著名的数据集。
            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括数据集格式和架构、标注、规模、训练模型专业方向等。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            重要注意事项：
            对数据集的内容进行详细规定，使用应该、不得、提供、应支持等规定性语言，不要泛泛而谈，对数据集大小、来源及标注等信息应该进行定量而非定性描述。
            user_prompt中的数据集简述提供该数据集的背景知识和功能，帮助你去撰写技术规范书。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            内容中不要出现*这个符号。
            在描述技术内容时，提供尽可能具体的细节，包括数据集结构。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            本次生成是以段落为单位,根据user_prompt去生成对应的段落,而不是整个文档,同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可。
            
            注意：
            1. 每个段落的目录分级，例如第一个段落引言，需要列出1.1, 1.2等目录，每个段落都需要目录。
            2. 应该使用规定性语言如应该、不得、提供、应支持等，以增强文档的权威性和规范性。
            3. 应该多使用英文术语和深度学习领域专业术语以提高规范性和权威性。
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            # 引言模板
            "introduction_template": PromptTemplate(
                input_variables=["system_message", "dataset_description", "dataset_name"],
                template="""
                {system_message}
                请为数据集{dataset_name}的技术规范书撰写引言段落，不少于1500字。
                参考文献应该保证准确性，不要出现张三，李四这种称谓。
                请包括1.1目的、1.2范围及1.3参考文献。以下是数据集简述：
                {dataset_description}
                
                """
                
            ),

            # 数据集概述模板
            "dataset_overview_template": PromptTemplate(
                input_variables=["system_message", "dataset_description", "dataset_name"],
                template="""
                {system_message}
                为数据集“{dataset_name}”撰写概述段落，不少于3000字。
                概述应包括数据集名称、来源、规模、范围、适用的模型及版本信息。
                使用应该、不得、提供、应支持等规定性语言，不要泛泛而谈，应该定量而非定性描述。
                需要列出2.1, 2.2等目录，每个段落都需要目录
                以下是数据集描述：
                {dataset_description}
                
                
                数据集规模的描述应该包括数据集的文件大小、数据条数、数据类型、数据格式等信息。
                简单示例：
                数据集规模：
                    数据集包含约10,00000条中药材记录。
                    总大小约为1.5 TB。
                    数据集包括多种文件格式，包括TXT、JSON、Excel和SQL文件。
                """
                ),

            # 数据集设计模板
            "dataset_design_template": PromptTemplate(
                input_variables=["system_message", "dataset_description", "dataset_name"],
                template="""
                {system_message}
                为数据集“{dataset_name}”撰写设计段落，不少于3000字。
                需要列出3.1, 3.2等目录，每个段落都需要目录
                使用应该、不得、提供、应支持等规定性语言，不要泛泛而谈，应该定量而非定性描述。
                描述数据采集、标注工具、标注标准、预处理方法、预标注及如何使用该数据集进行模型训练等，需要规定tensorflow等训练工具版本，包括：
                1. 数据来源及采集方法
                2. 标注工具及流程
                3. 数据预处理技术
                4. 预标注
                5. 模型训练, 模型主要以transformer和mamba等新模型为主。
                
                以下是数据集描述：
                {dataset_description}
                """
            ),

            # 数据集结构模板
            "dataset_structure_template": PromptTemplate(
                input_variables=["system_message", "dataset_description", "dataset_name"],
                template="""
                {system_message}
                描述数据集“{dataset_name}”的文件结构，不少于3000字。
                使用应该、不得、提供、应支持等规定性语言，不要泛泛而谈，应该定量而非定性描述。
                需要列出4.1, 4.2等目录，每个段落都需要目录
                包括数据文件格式、组织方式。
                以下是数据集描述：
                {dataset_description}
                应该包括两种数据集格式，一种是json数据集，一种是多轮对话数据集。
                注意以下例子是特定行业的数据集，仅供参考，请根据综述中实际的数据集所属专业和行业撰写。
                注意，在描述数据集格式时需要给出相应的示例，请注意以下示例只做格式参考，内容必须按照相应数据集的行业去更改。
                json数据集示例:
                [
                    {{
                        "id": 55034,
                        "category": "黑头",
                        "dialogues": [
                            {{
                                "role": "patient",
                                "content": "女 22岁 黑头多 毛孔粗大..."
                            }},
                            {{
                                "role": "doctor",
                                "content": "你好，你的皮肤中间T区..."
                            }}
                        ]
                    }},
                    {{
                        "id": 51845,
                        "category": "中耳炎",
                        "dialogues": [
                            {{
                                "role": "patient",
                                "content": "张医生，孩子今天去复查了..."
                            }},
                            {{
                                "role": "doctor",
                                "content": "您好，欢迎来问，感谢信任..."
                            }}
                        ]
                    }}
                ]
                
                多轮对话数据集:
                客户：你好。
                客服：您好。
                客户：我已经充值了话费，但是现在还是处于停机状态，这是怎么回事？
                客服：请您提供一下手机号码，我帮您查看一下具体情况。
                客户：好的，这是我的手机号码。
                客服：我查询到您今天下午三点多充值了五十元，款项已经到账。但是我需要再次确认一下，您稍等一下。
                客户：好的，我等待您的回复。
                客服：我这边查询到您的充值操作并未成功。
                客户：充值操作失败了？我是通过手机进行充值的。

                """
            ),

            # 数据集质量评估模板
            "dataset_quality_template": PromptTemplate(
                input_variables=["system_message", "dataset_description", "dataset_name"],
                template="""
                {system_message}
                为数据集“{dataset_name}”撰写质量评估段落，不少于3000字。
                质量指标应该以定量而非定性的方式给出，例如数据分布特性的方差，应该规定其具体数值。
                使用应该、不得、提供、应支持等规定性语言，不要泛泛而谈，应该定量而非定性描述。
                需要列出5.1, 5.2等目录，每个段落都需要目录
                以下是数据集描述：
                {dataset_description}
                描述数据集的质量指标，如:数据完整性, 数据准确性, 数据一致性, 数据覆盖度, 数据平衡性, 数据标注质量,数据噪声, 数据分布特性等。
                """
                )         
        }

    def generate_text(self, prompt_template, dataset_description, dataset_name, index, func_name = None, func_info = None, **kwargs):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == None:
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                dataset_description=dataset_description,
                dataset_name=dataset_name
            )
        else:
            # func_name 和 func_info 应该同时不为空，当撰写具体功能技术规范段落时
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                dataset_description=dataset_description,
                dataset_name=dataset_name,
                index = index,
                func_name=func_name,
                func_info=func_info
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text
    
    


    def generate_section(self, prompt_template, section_title, index, dataset_description, dataset_name, func_name = None, func_info = None, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            dataset_description=dataset_description,
            dataset_name=dataset_name,
            index = index,
            func_name=func_name,
            func_info=func_info
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string
    
    
    # 提取并保存采购报价单excel
    def save_func_names_json_to_excel(self, func_names_json, output_file_path):
        """
        将 func_names_json 数据处理成 Excel 文件并保存。

        参数:
        func_names_json (list): 包含功能模块和子功能描述的 JSON 数据。
        output_file_path (str): 输出 Excel 文件的路径。
        """
        # Prepare data for DataFrame
        data = []
        for module in func_names_json:
            module_name = module["module_name"]
            for sub_function in module["sub_functions"]:
                sub_function_name = sub_function["sub_function_name"]
                description = sub_function["description"]
                data.append({"功能模块": module_name, "子功能描述": sub_function_name, "具体内容描述": description})

        # Create DataFrame
        df = pd.DataFrame(data)

        # Create a workbook and worksheet
        wb = Workbook()
        ws = wb.active

        # Write DataFrame to worksheet
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=value)

        # Merge cells with the same "功能模块"
        prev_module_name = None
        start_row = 1
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
            current_module_name = row[0].value
            if current_module_name != prev_module_name:
                if prev_module_name is not None:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=row[0].row-1, end_column=1)
                prev_module_name = current_module_name
                start_row = row[0].row
        if prev_module_name is not None:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=ws.max_row, end_column=1)

        # Set alignment
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths with better handling for wide characters
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter  # Get column letter (e.g., 'A', 'B')
            for cell in col:
                if cell.value:
                    # Estimate character width: consider wide characters like Chinese as 2 units
                    value = str(cell.value)
                    adjusted_length = sum(2 if ord(char) > 127 else 1 for char in value)
                    max_length = max(max_length, adjusted_length)
            ws.column_dimensions[col_letter].width = max_length + 2  # Add padding

        # Save the Excel file
        wb.save(output_file_path)

    
    def format_func_info(self, func_info):
        module_name = func_info['module_name']
        sub_functions = func_info['sub_functions']
        
        # 初始化一个空字符串用于存储结果
        result = f"模块名称: {module_name}"
        
        # 遍历子功能并拼装字符串
        for sub_function in sub_functions:
            sub_function_name = sub_function['sub_function_name']
            description = sub_function['description']
            result += f"  子功能名称: {sub_function_name}"
            result += f"    描述: {description}"
        
        return result

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf8') as f:
            dataset_description = f.read()

        
        # 生成包含功能模块及其子功能描述的JSON
        func_names_json_prompt = f"""
            根据以下数据集简述, 从数据集制作的角度, 生成相应的步骤、子步骤和子步骤描述的JSON。
            一个数据集制作应有多个步骤，每个步骤包含多个子步骤，一个子步骤对应一个子步骤描述。
            至少生成4个步骤, 每个步骤至少包含四个及以上的子步骤。
            产品简述:
            {dataset_description}

            生成的JSON应符合以下格式：
            [
                {{
                    "module_name": "数据收集",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }},
                        {{
                            "sub_function_name": "子功能2",
                            "description": "具体内容描述2"
                        }}
                        // 可以添加更多子功能
                    ]
                }},
                {{
                    "module_name": " 数据清洗",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }}
                        // 可以添加更多子功能
                    ]
                }}
                // 可以添加更多模块
            ]
            """

        # 使用LLM生成JSON
        func_names_json_response = self.llm.predict(func_names_json_prompt)
        
        try:
            # Strip leading/trailing whitespace and newlines
            stripped_response = func_names_json_response.strip().strip('```json')
            
            # Parse the stripped response into JSON
            func_names_info = json.loads(stripped_response)
            
            # Use func_names_info as needed
            print("功能详情:", func_names_info)

        except json.JSONDecodeError as e:
            # Handle JSON decoding errors
            print(f"JSON decoding error: {e}")
            func_names_info = None  # Or set to a default value

        except Exception as e:
            # Handle any other unexpected errors
            print(f"An error occurred: {e}")
            func_names_info = None  # Or set to a default value
        try:
            self.save_func_names_json_to_excel(func_names_info, "./采购报价单.xlsx")
            print(f"采购报价单已保存")
        except json.JSONDecodeError as e:
            print(f"JSON 解码错误: {e}")
            func_names_json = None
        
        
        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction_template"]},
        {"title": "数据集概述", "prompt": self.prompts["dataset_overview_template"]},
        {"title": "数据集设计", "prompt": self.prompts["dataset_design_template"]},
        {"title": "数据集结构", "prompt": self.prompts["dataset_structure_template"]},
        {"title": "数据集质量评估", "prompt": self.prompts["dataset_quality_template"]},
        ]


        # Step 4: 生成数据集
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                index = index,
                dataset_description=dataset_description,
                dataset_name=self.dataset_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.dataset_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'dataset_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['dataset_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            dataset_name=platform["dataset_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['dataset_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['dataset_name']}: {e}")
            logger.error("Traceback details:", exc_info=True)
            

if __name__ == "__main__":
    # Define the list of platforms
    platforms = [
        {
            "docx_name": "法律数据集技术规范书.docx",
            "title": "法律数据集技术规范书",
            "dataset_name": "法律数据集"
        }
        
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
