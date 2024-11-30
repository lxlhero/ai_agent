import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
import docx
from typing import Any
import logging
from docx.shared import Pt
import traceback
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment



# 说明性质的技术规范书



# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        # 这里是整个产品的名称
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的产品描述和平台名称，撰写说明性的技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            重要注意事项：
            本次生成的是说明性质的技术规范书内容, 因此重点是说明产品的主要功能架构和技术特点, 是用于采购外来产品，因此不要过多涉及如何自研开发，这点非常重要！
            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括系统架构、模块功能、输入输出、处理逻辑和部署要求。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            
            user_prompt中的产品简述提供该产品的背景知识和功能，帮助你去撰写技术规范书。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            在描述技术内容时，提供尽可能具体的细节，例如关键组件、数据流和测试标准。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            本次生成是以段落为单位,根据user_prompt去生成对应的段落,而不是整个文档,同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可。
            如果是大模型或相关产品，需要规定GPU或TPU等硬件要求。
            如果是软件功能，请规定数据库设计，接口设计等。
            注意：
            1. 每个段落的目录分级，例如第一个段落引言，需要列出1.1, 1.2等目录，每个段落都需要目录。
            2. 应该使用客观规范的语言进行产品说明，以增强文档的权威性和规范性。
            3. 应该多使用英文术语以提高规范性和权威性。
            4. 如果产品简述中提及是说明性的，那么应该以说明产品的主要功能架构和技术特点为核心。
            
            语言风格示例:
            创建虚拟用户组
            LoadRunner Professional提供一个直观的用户界面，使管理员能够轻松创建新的虚拟用户组。每个虚拟用户组都具备唯一标识符和描述性标签，以便于后续的管理和识别。在创建虚拟用户组的过程中，系统支持对服务时间、语言能力和专业领域等参数的自定义配置，以满足多样化的测试需求。

            管理虚拟用户组
            LoadRunner Professional允许管理员对已创建的虚拟用户组进行编辑操作，包括但不限于修改名称、描述和配置参数。系统还提供虚拟用户组的启用和禁用功能，以便在不删除配置的情况下，临时停用某个虚拟用户组。此外，虚拟用户组的配置应支持版本控制机制，记录所有历史更改，并允许管理员恢复到任意历史配置状态。

            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{product_name}的技术规范书引言，简要介绍文档的目的和范围，不少于1000字。
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    生成的不是技术规范书的整体, 而是针对该模块的引言段落。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    1.1 目的
                    本文档旨在明确LoadRunner Professional的技术标准和规范，确保其在性能测试过程中的规范性和一致性。通过对LoadRunner Professional的功能、架构、接口和实施要求的详细阐述，旨在帮助用户理解和使用该产品，以满足性能测试的需求，并提升测试的效率和准确性。本文档将作为技术团队、测试团队及相关利益相关者的参考依据，确保各方对LoadRunner Professional的理解和应用达成一致。

                    1.2 范围
                        本技术规范书全面覆盖LoadRunner Professional的使用和操作所需的全部技术要求。具体包括但不限于：
                        系统架构描述，包括LoadRunner Professional的总体架构设计、模块功能和技术实现。
                        功能操作指南，涉及LoadRunner Professional的核心功能点以及附加功能的具体操作方式。
                        接口规范说明，明确LoadRunner Professional与其他测试工具或系统集成时的接口规范和数据交互格式。
                        安全性要求，确保LoadRunner Professional在数据存储、传输和处理过程中的安全性。
                        可靠性指标，保障LoadRunner Professional的高可用性、故障恢复能力和性能稳定性。
                        维护和支持指南，提供LoadRunner Professional的维护流程、升级策略和技术支持渠道。

                    1.3 参考文档
                        为确保本技术规范书的科学性和权威性，编制过程中参考了以下文档和标准：
                        ISO/IEC 25010:2011 系统和软件工程——系统和软件产品质量模型，提供了系统和软件质量的通用评价框架。
                        相关API接口文档，为LoadRunner Professional提供了与其他软件组件集成所需的接口信息和数据交互标准。
                        数据保护法规，包括GDPR、CCPA等，确保LoadRunner Professional在处理用户数据时遵守相应的法律法规。

                    1.4 定义和缩略语
                        为了确保本文档中的术语使用一致和清晰，以下列出了一些关键词汇及其定义：
                        LoadRunner Professional：一款性能测试工具，用于模拟真实用户行为，评估系统在高负载下的性能和稳定性。
                        Virtual User Generator：虚拟用户生成器，用于创建虚拟用户脚本，模拟真实用户的业务操作行为。
                        Controller：控制器，用于组织、驱动、管理和监控负载测试。
                        Analysis：分析器，用于查看、分析和比较性能测试结果。
                        API（Application Programming Interface）：应用程序编程接口，允许不同软件组件之间进行互相通信。
                        SDK（Software Development Kit）：软件开发工具包，提供了一套工具、指南和API，以便开发者创建软件应用程序。
                        HTTPS（Hypertext Transfer Protocol Secure）：安全超文本传输协议，通过对传输数据进行加密，提供了数据传输的安全性保障。
                        SSL/TLS（Secure Sockets Layer / Transport Layer Security）：安全套接层/传输层安全，两种用于在互联网上确保数据传输安全的协议。

                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{product_name}的系统概述，包括其主要功能和用途, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统概述段落，是整个技术规范书的第二段，第一段是引言，所以专注于系统概述本身，不需要涉及引言或系统架构等段落。
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                     2.1 系统功能

                    LoadRunner Professional旨在提供一套完整的性能测试解决方案，允许用户模拟真实世界中的用户负载，以测试应用程序的性能和稳定性。以下是该产品的主要功能：

                    2.1.1 虚拟用户生成与录制
                    Web/HTTP/HTML: 录制和回放基于浏览器的应用程序的HTTP请求。
                    SAP: 录制和回放SAP业务流程。
                    Java Vuser: 用于录制和回放Java应用程序。
                    .NET Vuser: 用于录制和回放.NET应用程序。

                    2.1.2 性能测试执行
                    Controller: 提供了一个中央控制台，用于设计、配置、运行和分析测试场景。
                    Load Generator: 在多台机器上运行虚拟用户，生成实际负载。

                    2.1.3 性能监控
                    RealTime Monitoring: 实时监控服务器和网络性能指标。
                    Transaction Tracing: 跟踪事务响应时间，帮助识别性能瓶颈。

                    2.1.4 结果分析与报告
                    Analysis: 提供了一个强大的分析工具，用于查看、分析和比较性能测试结果。
                    Reporting: 生成详细的性能报告，包括响应时间、吞吐量和资源利用率等指标。

                    2.2 用户特征

                    LoadRunner Professional面向以下用户群体：

                    性能测试工程师: 使用该工具进行应用程序的性能测试，确保系统在高负载下的稳定性和性能。
                    系统管理员: 利用LoadRunner Professional监控系统性能，进行故障排查和性能调优。
                    开发人员: 在软件开发过程中使用LoadRunner Professional进行性能测试，以提高软件质量。

                    2.3 运行环境

                    LoadRunner Professional的运行环境要求如下：

                    服务器端
                    操作系统: 支持Linux和Windows服务器操作系统。
                    硬件要求: 根据测试需求和虚拟用户数量，可能需要高性能的CPU、内存和存储设备。
                    网络: 确保足够的网络带宽和低延迟，以支持大规模的负载测试。

                    客户端
                    操作系统: 支持Windows、Linux和macOS操作系统。
                    浏览器: LoadRunner Professional的Web界面应与主流浏览器如Chrome、Firefox、Safari和Edge兼容。
                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{product_name}的系统架构, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统架构段落。
                    架构描述应包括产品的整体架构，设计逻辑和性能优势等，不需要涉及引言或系统概述等段落。
                    根据你对系统架构的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                     3.1 总体架构

                    LoadRunner Professional采用模块化、分布式的架构设计，旨在确保系统的可扩展性、可维护性和高可用性。系统的主要架构分为以下几个层次：

                    3.1.1 控制层
                    控制层是LoadRunner Professional的核心组件，负责管理和控制整个测试过程。该层主要包括：
                    Controller: 提供了一个中央控制台，用于设计、配置、运行和分析测试场景。
                    Load Generator: 在多台机器上运行虚拟用户，生成实际负载。

                    3.1.2 虚拟用户层
                    虚拟用户层负责模拟真实用户的行为，包括浏览网页、填写表单、提交事务等。该层主要包括：
                    Virtual User Generator: 用于创建和录制虚拟用户脚本，模拟用户在应用程序上的操作。

                    3.1.3 监控层
                    监控层负责实时监控测试过程中的系统性能和资源使用情况。该层主要包括：
                    RealTime Monitoring: 实时监控服务器和网络性能指标。
                    Transaction Tracing: 跟踪事务响应时间，帮助识别性能瓶颈。

                    3.2 系统模块划分

                    LoadRunner Professional根据功能需求被划分为多个模块，以实现职责的分离和功能的独立。

                    3.2.1 测试设计模块
                    Script Creation: 提供工具和界面，用于创建和编辑虚拟用户脚本。
                    Test Scenario Design: 允许用户设计和配置测试场景，包括虚拟用户的数量、行为和负载模式。

                    3.2.2 测试执行模块
                    Test Execution: 负责运行配置好的测试场景，控制虚拟用户的执行。
                    Load Balancing: 在多个Load Generator之间分配虚拟用户，以实现负载均衡。

                    3.2.3 测试分析模块
                    Results Analysis: 提供了一个强大的分析工具，用于查看、分析和比较性能测试结果。
                    Reporting: 生成详细的性能报告，包括响应时间、吞吐量和资源利用率等指标。

                    通过上述架构设计和模块划分，LoadRunner Professional能够为用户提供一个全面、高效的性能测试平台，满足各种复杂的性能测试需求。
                    """
                )
            ),
            
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index", "func_name", "func_info"],
                template=(
                    """
                    {system_message}
                    请详细描述功能“{func_name}”的技术规范段落, 不少于1800字。
                    这是该功能的具体子模块和描述，请参考：“{func_info}”
                    技术规范需要涵盖该功能的所有子模块技术要求， 并根据你对技术规范的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    技术规范应包括以下内容：
                    设计逻辑和架构：详细说明功能的设计逻辑和架构，以及其技术细节和优势。
                    功能定义：概述功能的核心目标、输入输出及其处理流程。
                    技术要求：列出该功能的关键技术需求，例如性能指标、可用性、响应时间等。
                    实现方案：提供高层次的实现细节，包括使用的技术框架、算法或工具, 具体到版本。
                    部署要求：说明功能的运行环境需求，如硬件、操作系统、依赖库等。
                    接入方式: 说明用户如何接入该功能。

                    在撰写时，请以清晰、结构化的方式逐点展开，确保内容详尽且专业。
                    该功能是产品“{product_name}”的重要组成部分，所以生成的不是技术规范书的整体，而是针对该功能的技术规范段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录。
                    
                    注意：
                    1. 应该使用客观规范的语言进行产品说明，以增强文档的权威性和规范性。
                    2. 应该多使用英文术语以提高规范性和权威性。
                    3. 不要出现段落标题，只写段落的内容即可。
                    
                    以下是产品简述：{product_description}
                    
                    以下是一个简短的示例，这里假设是第4段落，仅供参考，具体子目录要看{index}是多少, 注意，子目录一定要生成:
                     4.1 虚拟用户脚本创建

                    LoadRunner Professional提供了一套全面的虚拟用户脚本创建功能，旨在简化脚本的生成过程并确保其准确性和可维护性。该功能包括以下几个方面：

                    4.1.1 脚本录制

                    Virtual User Generator (VUGen): VUGen是LoadRunner Professional的核心组件之一，用于创建和录制虚拟用户脚本。用户可以通过VUGen录制用户在应用程序上的操作，生成相应的脚本。VUGen支持多种协议，如Web (HTTP/HTML)、SAP、Java、.NET等，确保能够覆盖各种应用程序类型。
                    录制模式: VUGen提供手动和自动两种录制模式，用户可以根据需要选择合适的模式进行脚本录制。手动模式下，用户可以精确控制录制的每一个步骤；自动模式下，VUGen会自动记录用户的操作。

                    4.1.2 脚本编辑

                    脚本编辑器: LoadRunner Professional内置了一个功能强大的脚本编辑器，支持语法高亮、代码提示和自动补全等功能，提高脚本编写的效率。用户可以在编辑器中对录制的脚本进行修改和优化。
                    调试工具: 为了确保脚本的正确性和稳定性，LoadRunner Professional提供了丰富的调试工具，包括断点设置、单步执行、变量监视等功能，帮助用户快速定位和修复脚本中的问题。

                    4.1.3 脚本参数化

                    参数化数据: LoadRunner Professional支持脚本参数化，允许用户使用外部数据源（如Excel、CSV文件）中的数据替换脚本中的常量值。这不仅提高了脚本的灵活性，还使得脚本能够模拟更多真实的用户行为。
                    参数集管理: 用户可以创建和管理多个参数集，每个参数集包含一组特定的参数值。在测试执行过程中，可以选择不同的参数集，以模拟不同的测试场景。

                    4.2 脚本执行与管理

                    LoadRunner Professional提供了一套完善的脚本执行和管理功能，确保测试过程的高效性和可控性。该功能包括以下几个方面：

                    4.2.1 测试场景配置

                    Controller: LoadRunner Professional的Controller用于组织、驱动、管理和监控负载测试。用户可以在Controller中创建测试场景，定义虚拟用户的数量、行为和负载模式。
                    调度策略: Controller支持多种调度策略，如逐步增加负载、保持恒定负载和逐步减少负载，以满足不同的测试需求。

                    4.2.2 负载生成

                    Load Generator: LoadRunner Professional的Load Generator负责在多台机器上运行虚拟用户，生成实际负载。Load Generator可以部署在不同的物理机、虚拟机或云平台上，以实现灵活的负载生成。
                    负载均衡: LoadRunner Professional支持负载均衡功能，确保虚拟用户在多台Load Generator之间均匀分布，避免单点过载。

                    4.2.3 测试执行监控

                    实时监控: LoadRunner Professional提供实时监控功能，允许用户在测试执行过程中实时查看系统性能指标，如CPU利用率、内存使用率、网络带宽等。
                    报警机制: 当系统性能指标超过预设阈值时，LoadRunner Professional会自动触发报警，提醒用户及时处理潜在问题。

                    4.3 脚本分析与报告

                    LoadRunner Professional提供了一套全面的脚本分析功能，帮助用户深入理解测试结果并进行优化。该功能包括以下几个方面：

                    4.3.1 性能数据分析

                    Analysis: LoadRunner Professional的Analysis工具用于分析测试结果，提供详细的性能数据，如响应时间、事务吞吐量、错误率等。
                    图表展示: Analysis工具支持多种图表类型，如折线图、柱状图、散点图等，帮助用户直观地理解性能数据。

                    4.3.2 报告生成

                    报告模板: LoadRunner Professional提供多种预定义的报告模板，用户可以根据需要选择合适的模板生成测试报告。
                    自定义报告: 用户也可以根据具体需求自定义报告内容和格式，生成符合特定要求的测试报告。

                    4.3.3 结果比较

                    历史数据对比: LoadRunner Professional支持历史数据对比功能，用户可以将当前测试结果与历史数据进行对比，分析系统性能的变化趋势。
                    基准测试: 用户可以创建基准测试，记录系统在不同负载下的性能表现，以便后续对比和分析。

                    通过上述虚拟用户脚本创建、执行与管理以及分析与报告功能，LoadRunner Professional能够为用户提供一个强大且灵活的性能测试平台，确保性能测试脚本的创建、维护和执行的准确性和高效性。
                """
                    
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index"],
                template=(
                    """
                    {system_message}
                    请描述{product_name}的维护和支持策略，包括维护计划和支持渠道,不少于400字。
                    注意是说明性质的，因此维护和支持应以产品的官方文档为准。
                    维护策略应说明{product_name}的维护计划和流程。
                    生成的不是技术规范书的整体, 而是针对该模块的维护与支持段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录
                    以下是产品简述：{product_description}
                    """
                )
            )
            
            
        }

    def generate_text(self, prompt_template, product_description, product_name, index, func_name = None, func_info = None, **kwargs):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == None:
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name
            )
        else:
            # func_name 和 func_info 应该同时不为空，当撰写具体功能技术规范段落时
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name,
                index = index,
                func_name=func_name,
                func_info=func_info
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text
    
    


    def generate_section(self, prompt_template, section_title, index, product_description, product_name, func_name = None, func_info = None, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            product_description=product_description,
            product_name=product_name,
            index = index,
            func_name=func_name,
            func_info=func_info
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string
    
    
    # 提取并保存采购报价单excel
    def save_func_names_json_to_excel(self, func_names_json, output_file_path):
        """
        将 func_names_json 数据处理成 Excel 文件并保存。

        参数:
        func_names_json (list): 包含功能模块和子功能描述的 JSON 数据。
        output_file_path (str): 输出 Excel 文件的路径。
        """
        # Prepare data for DataFrame
        data = []
        for module in func_names_json:
            module_name = module["module_name"]
            for sub_function in module["sub_functions"]:
                sub_function_name = sub_function["sub_function_name"]
                description = sub_function["description"]
                data.append({"功能模块": module_name, "子功能描述": sub_function_name, "具体内容描述": description})

        # Create DataFrame
        df = pd.DataFrame(data)

        # Create a workbook and worksheet
        wb = Workbook()
        ws = wb.active

        # Write DataFrame to worksheet
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=value)

        # Merge cells with the same "功能模块"
        prev_module_name = None
        start_row = 1
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
            current_module_name = row[0].value
            if current_module_name != prev_module_name:
                if prev_module_name is not None:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=row[0].row-1, end_column=1)
                prev_module_name = current_module_name
                start_row = row[0].row
        if prev_module_name is not None:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=ws.max_row, end_column=1)

        # Set alignment
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths with better handling for wide characters
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter  # Get column letter (e.g., 'A', 'B')
            for cell in col:
                if cell.value:
                    # Estimate character width: consider wide characters like Chinese as 2 units
                    value = str(cell.value)
                    adjusted_length = sum(2 if ord(char) > 127 else 1 for char in value)
                    max_length = max(max_length, adjusted_length)
            ws.column_dimensions[col_letter].width = max_length + 2  # Add padding

        # Save the Excel file
        wb.save(output_file_path)

    
    def format_func_info(self, func_info):
        module_name = func_info['module_name']
        sub_functions = func_info['sub_functions']
        
        # 初始化一个空字符串用于存储结果
        result = f"模块名称: {module_name}"
        
        # 遍历子功能并拼装字符串
        for sub_function in sub_functions:
            sub_function_name = sub_function['sub_function_name']
            description = sub_function['description']
            result += f"  子功能名称: {sub_function_name}"
            result += f"    描述: {description}"
        
        return result

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf8') as f:
            product_description = f.read()

        
        # 生成包含功能模块及其子功能描述的JSON
        func_names_json_prompt = f"""
            根据以下产品简述，生成相应的功能模块、子模块和子功能描述的JSON。
            一个产品应有多个功能模块，每个功能模块包含多个子模块，一个子模块对应一个子模块描述。
            至少生成4个功能模块, 每个功能模块至少包含四个及以上的子功能模块。
            产品简述:
            {product_description}

            生成的JSON应符合以下格式：
            [
                {{
                    "module_name": "模块1",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }},
                        {{
                            "sub_function_name": "子功能2",
                            "description": "具体内容描述2"
                        }}
                        // 可以添加更多子功能
                    ]
                }},
                {{
                    "module_name": "模块2",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }}
                        // 可以添加更多子功能
                    ]
                }}
                // 可以添加更多模块
            ]
            """

        # 使用LLM生成JSON
        func_names_json_response = self.llm.predict(func_names_json_prompt)
        
        try:
            # Strip leading/trailing whitespace and newlines
            stripped_response = func_names_json_response.strip().strip('```json')
            
            # Parse the stripped response into JSON
            func_names_info = json.loads(stripped_response)
            
            # Use func_names_info as needed
            print("功能详情:", func_names_info)

        except json.JSONDecodeError as e:
            # Handle JSON decoding errors
            print(f"JSON decoding error: {e}")
            func_names_info = None  # Or set to a default value

        except Exception as e:
            # Handle any other unexpected errors
            print(f"An error occurred: {e}")
            func_names_info = None  # Or set to a default value
        try:
            self.save_func_names_json_to_excel(func_names_info, "./采购报价单.xlsx")
            print(f"采购报价单已保存")
        except json.JSONDecodeError as e:
            print(f"JSON 解码错误: {e}")
            func_names_json = None
        
        func_names = [item['module_name'] for item in func_names_info]
        
        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction"]},
        {"title": "系统概述", "prompt": self.prompts["system_overview"]},
        {"title": "系统架构", "prompt": self.prompts["system_architecture"]},
        ]

        # Step 4: 生成前三个部分
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                index = index,
                product_description=product_description,
                product_name=self.product_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

        
        # Step 6: 为每个功能模块生成技术规范
        for index, func_name in enumerate(func_names[:], start=4):  # 从第 4 章开始
            logger.info(f"生成功能技术规范: {func_name}")
            # 查找与 func_name 匹配的模块
            func_info = next((module for module in func_names_info if module['module_name'] == func_name), None)
            
            func_info_str = self.format_func_info(func_info)
            
            if func_info_str is None:
                logger.error(f"未找到功能名称: {func_name}")
                continue
            try:
                section_paragraph = self.generate_section(
                    prompt_template=self.prompts["technical_specification"],
                    section_title=f"{index}. {func_name}技术规范",
                    index = index,
                    product_description=product_description,
                    product_name=self.product_name,
                    func_name=func_name,
                    func_info=func_info_str
                )
                # 清理多余字符并写入文档
                cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
                document.add_paragraph(cleaned_paragraph)
                
            except KeyError as e:
                logger.error(f"KeyError: Missing key '{e.args[0]}' in generate_text method")
                logger.error("Traceback details:", exc_info=True)
            except Exception as e:
                logger.error(f"Error generating technical specification for {product_name}: {e}")
                logger.error("Traceback details:", exc_info=True)

        
        # 添加维护与支持部分
        try:
            maintenance_index = len(func_names) + 4  
            maintenance_paragraph = self.generate_section(
                prompt_template=self.prompts["maintenance_support"],
                section_title=f"{maintenance_index}. 维护与支持",
                index = maintenance_index,
                product_description=product_description,
                product_name=self.product_name,
                func_name=func_name,
                func_info=func_info_str
            )
            cleaned_paragraph = maintenance_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)
        except Exception as e:
            logger.error(f"Error generating maintenance and support section: {e}")
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")
            logger.error("Traceback details:", exc_info=True)
            

if __name__ == "__main__":
    # Define the list of platforms
    platforms = [
        {
            "docx_name": "LoadRunner Professional技术规范书.docx",
            "title": "LoadRunner Professional技术规范书",
            "product_name": "LoadRunner Professional"
        }
        
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
