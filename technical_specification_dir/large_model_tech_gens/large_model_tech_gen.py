import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.chains import LLMChain
import docx
from typing import Any
import logging
from langchain.chains import LLMChain
from docx.shared import Pt


# 大模型模块技术规范书生成器
# 一次性生成以下所有模块的技术规范书
# 模块: 
# 数据预处理平台， 
# 模型训练平台，  
# 模型保存与部署平台， 
# 模型评估平台， 
# 模型推理平台


# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个文档编写者，你的任务是根据产品概述和平台名称撰写技术规范书的段落,
            此平台是该产品的组成部分，例如数据预处理平台、模型推理平台等，这是为该平台单独写技术规范书，
            要专注于这个平台本身，而不是整个产品的全部功能，不要过多设计其他平台的功能。
            技术规范书的要点是明确目的和范围、结构化和组织良好、详细和具体、准确性和可靠性，
            注意你编写的是大模型产品的功能性平台的技术规范书，例如数据预处理平台、模型推理平台等，所以要更贴合该平台的规范和特点
            不要去过多涉及该平台之外的功能，例如模型训练平台，就不要去写关于数据预处理或模型推理的内容，
            同时针对整体的产品概述，撰写的内容应该要贴合产品的特点，例如产品是面向工业界的，那么撰写的内容就要贴合工业界的规范和特点
            注意每个段落不要在开头出现段落title,比如撰写引言时，不要出现引言二字。
            以下是具体的需求:
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{module_name}的技术规范书引言，简要介绍文档的目的和范围，不少于500字。
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{module_name}的系统概述，该平台是下述产品简述的一部分，包括其主要功能和用途, 不少于600字。
                    产品简述：{product_description}
                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{module_name}的架构，该平台是下述产品简述的一部分，不少于800字。
                    架构描述应包括数据流、处理步骤和关键组件。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "function_module_description": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{module_name}的功能模块，该平台是下述产品简述的一部分，不少于800字。
                    功能模块描述应包括每个模块的作用、输入输出和处理逻辑。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请描述{module_name}的技术规范，该平台是下述产品简述的一部分，不少于800字。
                    技术规范应定义模块的规范、部署要求等。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "testing_validation": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请编写测试和验证部分的文档，包括测试计划、验证和验收标准，该平台是下述产品简述的一部分，不少于400字。
                    测试计划应描述{module_name}的测试策略,不要写测试用例。
                    验证和验收标准应说明{module_name}的验证和验收标准。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请描述{module_name}的维护和支持策略，包括维护计划和支持渠道，该平台是下述产品简述的一部分，不少于400字。
                    维护策略应说明{module_name}的维护计划和流程。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "conclusion": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请编写结论部分，简要总结{module_name}的主要内容和要点，该平台是下述产品简述的一部分，不少于500字。
                    总结应概述{module_name}的主要内容和要点。
                    下一步行动应说明文档发布后的下一步行动计划。
                    以下是产品简述：{product_description}
                    """
                )
            ),
            "appendix": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请编写附录部分，包括术语表和参考资料。
                    术语表应提供文档中使用的专业术语的解释。
                    参考资料应列出相关的标准、法规和其他参考文献。
                    """
                )
            )
        }

    def generate_text(self, prompt_template, **kwargs):
        """
            Generate text using LLMChain with the provided prompt template.
        """
        chain = LLMChain(prompt=prompt_template, llm=self.llm)
        generated_text = chain.run(**kwargs)
        
        logger.info("Generated Text: %s", generated_text)
        return generated_text


    def generate_section(self, prompt_template, section_title, product_description, module_name, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        system_message = self.system_message
        
        # Generate text content for the section
        section_content = self.generate_text(prompt_template, product_description=product_description, module_name=module_name, system_message = system_message, **kwargs)
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()

        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
            {"title": "引言", "prompt": self.prompts["introduction"]},
            {"title": "系统概述", "prompt": self.prompts["system_overview"]},
            {"title": "系统架构", "prompt": self.prompts["system_architecture"]},
            {"title": "功能模块说明", "prompt": self.prompts["function_module_description"]},
            {"title": "技术规范", "prompt": self.prompts["technical_specification"]},
            {"title": "测试与验证", "prompt": self.prompts["testing_validation"]},
            {"title": "维护和支持", "prompt": self.prompts["maintenance_support"]},
            {"title": "结论", "prompt": self.prompts["conclusion"]},
            {"title": "附录", "prompt": self.prompts["appendix"]}
        ]

        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(section["prompt"], f"{index}. {section['title']}", product_description, self.product_name)

            # 清理多余的 * 和 #
            cleaned_paragraph = section_paragraph.replace('*', '').replace('#', '')
            
            # 创建一个新的段落并设置字体样式
            document.add_paragraph(cleaned_paragraph)
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")

if __name__ == "__main__":
    # Define the list of platforms
    product_name = "MoE算法V1.0"
    platforms = [
        {
            "docx_name": "MoE专家网络设计规范书.docx",
            "title": f"{product_name} - 专家网络设计规范书",
            "product_name": "专家网络设计"
        },
        {
            "docx_name": "MoE门控机制规范书.docx",
            "title": f"{product_name} - 门控机制规范书",
            "product_name": "门控机制"
        },
        {
            "docx_name": "MoE训练策略规范书.docx",
            "title": f"{product_name} - 训练策略规范书",
            "product_name": "训练策略"
        },
        {
            "docx_name": "MoE模型评估与优化规范书.docx",
            "title": f"{product_name} - 模型评估与优化规范书",
            "product_name": "模型评估与优化"
        },
        {
            "docx_name": "MoE部署与应用规范书.docx",
            "title": f"{product_name} - 部署与应用规范书",
            "product_name": "部署与应用"
        }
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
