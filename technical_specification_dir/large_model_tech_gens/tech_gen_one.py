import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
import docx
from typing import Any
import logging
from docx.shared import Pt
import traceback
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment



# 大模型模块技术规范书生成器
# 根据产品简述生成技术规范书



# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        # 这里是整个产品的名称
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的产品描述和平台名称，撰写技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括系统架构、模块功能、输入输出、处理逻辑和部署要求。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            重要注意事项：
            user_prompt中的产品简述提供该产品的背景知识和功能，帮助你去撰写技术规范书。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            在描述技术内容时，提供尽可能具体的细节，例如关键组件、数据流和测试标准。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            本次生成是以段落为单位,根据user_prompt去生成对应的段落,而不是整个文档,同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可。
            如果是大模型或相关产品，需要规定GPU或TPU等硬件要求。
            如果是软件功能，请规定数据库设计，接口设计等。
            注意每个段落的目录分级，例如第一个段落引言，需要列出1.1, 1.2等目录，每个段落都需要目录。
            
            多使用规定性语言如应该、不得、提供、应支持等，以增强文档的权威性和规范性。
            语言风格示例:
            创建技能组
             提供一个用户友好的界面，允许管理员创建新的技能组。
             管理员可以为每个技能组指定一个唯一的名称和描述，以便于识别和管理。
             技能组的创建过程应支持自定义配置，如服务时间、语言能力、专业领域等。
            管理技能组
             管理员可以编辑已存在的技能组信息，包括修改名称、描述和配置参数。
             提供技能组的启用和禁用功能，允许临时停用某个技能组而不删除其配置。
             技能组配置应支持版本控制，记录历史更改并允许恢复到之前的配置。

            
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{product_name}的技术规范书引言，简要介绍文档的目的和范围，不少于1000字。
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    生成的不是技术规范书的整体, 而是针对该模块的引言段落。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    1.1 目的
                    本文档的编制目的是为了定义智能客服系统的技术标准和规范，确保系统开发、部署及运维过程的规范性和一致性。通过对系统的功能、性能、接口、设计和实施要求的详细规定，旨在保障系统能够满足业务需求，同时提升系统的稳定性、安全性和用户体验。本文档将作为项目开发团队、测试团队、运维团队以及相关利益相关者的参考依据，确保各方对系统的理解和期望达成一致。
                    1.2 范围
                    本技术规范书全面覆盖智能客服系统的构建和运行所需的全部技术要求。具体包括但不限于：
                     系统架构设计，包括系统的总体设计思路、模块划分和数据流向。
                     功能实现细节，涉及系统的核心功能点以及附加功能的具体实现方式。
                     接口标准规定，明确系统与外部系统集成时的接口规范和数据交互格式。
                     安全性要求，确保系统在数据存储、传输和处理过程中的安全性。
                     可靠性指标，保障系统的高可用性、故障恢复能力和性能稳定性。
                     维护和支持指南，提供系统的维护流程、升级策略和技术支持渠道。
                    1.3 参考文档
                    为确保本技术规范书的科学性和权威性，编制过程中参考了以下文档和标准：
                     ISO/IEC 25010:2011 系统和软件工程——系统和软件产品质量模型，提供了系统和软件质量的通用评价框架。
                     相关API接口文档，为系统提供了与外部服务和系统集成所需的接口信息和数据交互标准。
                     数据保护法规，包括GDPR、CCPA等，确保系统在处理用户数据时遵守相应的法律法规。
                    1.4 定义和缩略语
                    为了确保本文档中的术语使用一致和清晰，以下列出了一些关键词汇及其定义：
                     CRM（Customer Relationship Management）：客户关系管理系统，用于帮助企业管理与客户的互动信息。
                     API（Application Programming Interface）：应用程序编程接口，允许不同软件组件之间进行互相通信。
                     SDK（Software Development Kit）：软件开发工具包，提供了一套工具、指南和API，以便开发者创建软件应用程序。
                     UI（User Interface）：用户界面，指系统与用户交互的界面部分，包括布局、视觉元素和交互逻辑。
                     UX（User Experience）：用户体验，指用户在使用系统过程中的总体感受和体验。
                     HTTPS（Hypertext Transfer Protocol Secure）：安全超文本传输协议，通过对传输数据进行加密，提供了数据传输的安全性保障。
                     SSL/TLS（Secure Sockets Layer / Transport Layer Security）：安全套接层/传输层安全，两种用于在互联网上确保数据传输安全的协议。
                     NLP（Natural Language Processing）：自然语言处理，指使计算机能够理解和处理人类语言的技术。

                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{product_name}的系统概述，包括其主要功能和用途, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统概述段落，是整个技术规范书的第二段，第一段是引言，所以专注于系统概述本身，不需要涉及引言或系统架构等段落。
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    2.1 系统功能
                    智能客服系统旨在整合多种通信渠道和客服资源，提供高效、便捷的客户服务体验。以下是系统的主要功能：
                     2.1.1 多渠道接入：
                     桌面网站：通过提供可嵌入网页的聊天插件，使用户能够直接在网站上与客服交流。
                     移动网站：确保聊天插件能够在移动设备上自动适配屏幕大小，保持良好的用户体验。
                     App：提供适用于iOS和Android平台的SDK，使得App内可以无缝接入客服功能。
                     微信：集成微信公众号和小程序的客服接口，直接在微信环境中提供客服服务。
                     微博：通过微博开放平台实现客服功能，与用户进行互动。
                     短信：接入短信服务平台，为无法使用即时通信的用户提供客服支持。
                     2.1.2 接待方式：
                     人工客服：系统为人工客服提供完整的操作界面，包括即时消息处理、用户信息查看、历史记录访问等。
                     机器人客服：集成自然语言处理技术，实现对用户咨询的自动响应，并提供转接人工客服的选项。
                     技能组管理：管理员可以创建不同的技能组，根据客服人员的专长将其分配至相应组别，系统根据用户咨询内容自动分配至对应技能组。
                     富媒体沟通：
                     文字：支持发送接收即时文本消息。
                     表情：集成表情包，丰富交流情感表达。
                     图片：允许发送图片文件，并在聊天窗口中直接预览。
                     富文本：支持发送包含格式的文本消息，如加粗、斜体、列表等。
                    
                    2.2 用户特征
                    智能客服系统面向以下用户群体：
                     客服人员：使用系统接待访客，处理咨询和问题，需要系统提供稳定、高效的操作界面。
                     系统管理员：负责系统的设置、维护和升级，需要管理工具来监控系统状态和性能。
                     访客：通过不同渠道接触系统，寻求即时的帮助和信息，期待快速、准确的服务。
                    2.3 运行环境
                    智能客服系统的运行环境要求如下：
                     服务器端：
                     操作系统：支持Linux和Windows服务器操作系统。
                     部署：支持物理服务器、虚拟化环境及云服务平台部署。
                     数据库：兼容主流数据库系统，如MySQL、PostgreSQL等。
                     客户端：
                     浏览器：确保与主流浏览器如Chrome、Firefox、Safari和Edge兼容。
                     移动操作系统：支持iOS和Android操作系统，确保App内客服功能的稳定运行。

                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{product_name}的系统架构, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统架构段落。
                    架构描述应包括数据流、处理步骤和关键组件。
                    根据你对系统架构的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    3.1 总体架构
                    智能客服系统将采用一种模块化、分布式的架构设计，以确保系统的可扩展性、可维护性和高可用性。系统的主要架构分为以下几个层次：
                    3.1.1 前端展示层
                    前端展示层负责提供用户界面(UI)，使用户能够通过图形界面与系统进行交互。该层主要包括：
                     客户端UI：为不同终端（桌面、移动、App等）提供定制化的用户界面。
                     Web UI：为桌面和移动浏览器提供响应式设计的网页界面。
                     第三方平台UI：为微信、微博等第三方平台集成的客服界面。
                    3.1.2 业务逻辑层
                    业务逻辑层是系统的核心，负责处理客服业务流程和用户交互逻辑。包括但不限于：
                     会话管理：处理用户与客服的会话，包括会话的创建、维持、结束和转接。
                     技能组路由：根据用户的问题和预设的规则，将用户的咨询定向到合适的技能组。
                     消息处理：处理发送和接收的消息，包括文本、图片、视频等多种格式。
                    
                    3.2 系统模块划分
                    系统根据功能需求被划分为多个模块，以实现职责的分离和功能的独立。
                    3.2.1 用户管理模块
                     用户认证：负责用户登录认证过程，包括密码校验、会话管理等。
                     用户权限：管理用户的权限设置，确保用户只能访问授权的资源。
                     用户信息：维护用户的基本信息，包括联系方式、偏好设置等。
                    3.2.2 消息处理模块
                     消息队列：采用消息队列进行消息的收发，保证消息的可靠传输。
                     消息格式化：将消息转换为标准格式，以便在不同的通信渠道中传输。
                     消息记录：记录所有的消息历史，以供查询和分析。
                    3.2.3 数据分析模块
                     访问分析：分析用户的访问行为和路径，生成访问报告。
                     会话分析：统计会话的相关数据，如持续时间、转化率等。
                     用户满意度：收集和分析用户满意度调查的结果。
                    
                    """
                )
            ),
            
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index", "func_name", "func_info"],
                template=(
                    """
                    {system_message}
                    请详细描述功能“{func_name}”的技术规范段落, 不少于1800字。
                    这是该功能的具体子模块和描述，请参考：“{func_info}”
                    技术规范需要涵盖该功能的所有子模块技术要求， 并根据你对技术规范的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    技术规范应包括以下内容：
                    功能定义：概述功能的核心目标、输入输出及其处理流程。
                    技术要求：列出该功能的关键技术需求，例如性能指标、可用性、响应时间等。
                    实现方案：提供高层次的实现细节，包括使用的技术框架、算法或工具, 具体到版本。
                    数据库设计和接口设计。
                    部署要求：说明功能的运行环境需求，如硬件、操作系统、依赖库等。
                    接入方式: 说明用户如何接入该功能，如浏览器接入，微信接入等。

                    在撰写时，请以清晰、结构化的方式逐点展开，确保内容详尽且专业。
                    该功能是产品“{product_name}”的重要组成部分，所以生成的不是技术规范书的整体，而是针对该功能的技术规范段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录。
                    
                    不要出现段落标题，只写段落的内容即可。
                    以下是产品简述：{product_description}
                    
                    以下是一个示例，这里假设是第4段落，仅供参考，具体子目录要看{index}是多少, 注意，子目录一定要生成:
                    4.1 桌面网站接入
                    4.1.1 JavaScript SDK
                     提供JavaScript SDK以供集成到企业桌面网站中，该SDK应包含以下功能：
                     用户身份验证：确保用户与系统之间的通信是安全的。
                     会话管理：允许创建、维护和结束用户会话。
                     消息发送和接收：支持文本、图片、文件等多种消息类型的发送和接收。
                     事件监听：能够响应用户交互事件，如点击、输入等。
                     UI组件：提供标准的聊天窗口UI组件，支持自定义样式以符合企业品牌形象。
                    4.1.2 通讯机制
                     系统应支持WebSocket协议，以实现实时、双向的通信。
                     对于不支持WebSocket的环境，系统应提供长轮询的备选方案，以保证消息的可靠传输。
                     所有通讯应通过HTTPS进行加密，以确保数据传输的安全性。
                    4.2 移动网站接入
                    4.2.1 自适应设计
                     移动端接入应使用响应式设计，确保UI组件能够根据不同屏幕尺寸进行适配。
                     SDK应提供与桌面网站接入相同的功能，并优化触摸操作和网络条件变化的响应。
                    4.3 数据库表设计
                     用户表设计：
                         `user_id` INT PRIMARY KEY AUTO_INCREMENT
                         `username` VARCHAR(255) UNIQUE NOT NULL
                         `password_hash` VARCHAR(255) NOT NULL
                         `role` ENUM('admin', 'agent', 'customer') NOT NULL

                    """
                    
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index"],
                template=(
                    """
                    {system_message}
                    请描述{product_name}的维护和支持策略，包括维护计划和支持渠道,不少于400字。
                    维护策略应说明{product_name}的维护计划和流程。
                    生成的不是技术规范书的整体, 而是针对该模块的维护与支持段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录
                    以下是产品简述：{product_description}
                    """
                )
            )
            
            
        }

    def generate_text(self, prompt_template, product_description, product_name, index, func_name = None, func_info = None, **kwargs):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == None:
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name
            )
        else:
            # func_name 和 func_info 应该同时不为空，当撰写具体功能技术规范段落时
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name,
                index = index,
                func_name=func_name,
                func_info=func_info
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text
    
    


    def generate_section(self, prompt_template, section_title, index, product_description, product_name, func_name = None, func_info = None, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            product_description=product_description,
            product_name=product_name,
            index = index,
            func_name=func_name,
            func_info=func_info
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string
    
    
    # 提取并保存采购报价单excel
    def save_func_names_json_to_excel(self, func_names_json, output_file_path):
        """
        将 func_names_json 数据处理成 Excel 文件并保存。

        参数:
        func_names_json (list): 包含功能模块和子功能描述的 JSON 数据。
        output_file_path (str): 输出 Excel 文件的路径。
        """
        # Prepare data for DataFrame
        data = []
        for module in func_names_json:
            module_name = module["module_name"]
            for sub_function in module["sub_functions"]:
                sub_function_name = sub_function["sub_function_name"]
                description = sub_function["description"]
                data.append({"功能模块": module_name, "子功能描述": sub_function_name, "具体内容描述": description})

        # Create DataFrame
        df = pd.DataFrame(data)

        # Create a workbook and worksheet
        wb = Workbook()
        ws = wb.active

        # Write DataFrame to worksheet
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=value)

        # Merge cells with the same "功能模块"
        prev_module_name = None
        start_row = 1
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
            current_module_name = row[0].value
            if current_module_name != prev_module_name:
                if prev_module_name is not None:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=row[0].row-1, end_column=1)
                prev_module_name = current_module_name
                start_row = row[0].row
        if prev_module_name is not None:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=ws.max_row, end_column=1)

        # Set alignment
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths with better handling for wide characters
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter  # Get column letter (e.g., 'A', 'B')
            for cell in col:
                if cell.value:
                    # Estimate character width: consider wide characters like Chinese as 2 units
                    value = str(cell.value)
                    adjusted_length = sum(2 if ord(char) > 127 else 1 for char in value)
                    max_length = max(max_length, adjusted_length)
            ws.column_dimensions[col_letter].width = max_length + 2  # Add padding

        # Save the Excel file
        wb.save(output_file_path)

    
    def format_func_info(self, func_info):
        module_name = func_info['module_name']
        sub_functions = func_info['sub_functions']
        
        # 初始化一个空字符串用于存储结果
        result = f"模块名称: {module_name}"
        
        # 遍历子功能并拼装字符串
        for sub_function in sub_functions:
            sub_function_name = sub_function['sub_function_name']
            description = sub_function['description']
            result += f"  子功能名称: {sub_function_name}"
            result += f"    描述: {description}"
        
        return result

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf8') as f:
            product_description = f.read()

        
        # 生成包含功能模块及其子功能描述的JSON
        func_names_json_prompt = f"""
            根据以下产品简述，生成相应的功能模块、子模块和子功能描述的JSON。
            一个产品应有多个功能模块，每个功能模块包含多个子模块，一个子模块对应一个子模块描述。
            至少生成4个功能模块, 每个功能模块至少包含四个及以上的子功能模块。
            产品简述:
            {product_description}

            生成的JSON应符合以下格式：
            [
                {{
                    "module_name": "模块1",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }},
                        {{
                            "sub_function_name": "子功能2",
                            "description": "具体内容描述2"
                        }}
                        // 可以添加更多子功能
                    ]
                }},
                {{
                    "module_name": "模块2",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }}
                        // 可以添加更多子功能
                    ]
                }}
                // 可以添加更多模块
            ]
            """

        # 使用LLM生成JSON
        func_names_json_response = self.llm.predict(func_names_json_prompt)
        
        try:
            # Strip leading/trailing whitespace and newlines
            stripped_response = func_names_json_response.strip().strip('```json')
            
            # Parse the stripped response into JSON
            func_names_info = json.loads(stripped_response)
            
            # Use func_names_info as needed
            print("功能详情:", func_names_info)

        except json.JSONDecodeError as e:
            # Handle JSON decoding errors
            print(f"JSON decoding error: {e}")
            func_names_info = None  # Or set to a default value

        except Exception as e:
            # Handle any other unexpected errors
            print(f"An error occurred: {e}")
            func_names_info = None  # Or set to a default value
        try:
            self.save_func_names_json_to_excel(func_names_info, "./采购报价单.xlsx")
            print(f"采购报价单已保存")
        except json.JSONDecodeError as e:
            print(f"JSON 解码错误: {e}")
            func_names_json = None
        
        func_names = [item['module_name'] for item in func_names_info]
        
        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction"]},
        {"title": "系统概述", "prompt": self.prompts["system_overview"]},
        {"title": "系统架构", "prompt": self.prompts["system_architecture"]},
        ]

        # Step 4: 生成前三个部分
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                index = index,
                product_description=product_description,
                product_name=self.product_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

        
        # Step 6: 为每个功能模块生成技术规范
        for index, func_name in enumerate(func_names[:], start=4):  # 从第 4 章开始
            logger.info(f"生成功能技术规范: {func_name}")
            # 查找与 func_name 匹配的模块
            func_info = next((module for module in func_names_info if module['module_name'] == func_name), None)
            
            func_info_str = self.format_func_info(func_info)
            
            if func_info_str is None:
                logger.error(f"未找到功能名称: {func_name}")
                continue
            try:
                section_paragraph = self.generate_section(
                    prompt_template=self.prompts["technical_specification"],
                    section_title=f"{index}. {func_name}技术规范",
                    index = index,
                    product_description=product_description,
                    product_name=self.product_name,
                    func_name=func_name,
                    func_info=func_info_str
                )
                # 清理多余字符并写入文档
                cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
                document.add_paragraph(cleaned_paragraph)
                
            except KeyError as e:
                logger.error(f"KeyError: Missing key '{e.args[0]}' in generate_text method")
                logger.error("Traceback details:", exc_info=True)
            except Exception as e:
                logger.error(f"Error generating technical specification for {product_name}: {e}")
                logger.error("Traceback details:", exc_info=True)

        
        # 添加维护与支持部分
        try:
            maintenance_index = len(func_names) + 4  
            maintenance_paragraph = self.generate_section(
                prompt_template=self.prompts["maintenance_support"],
                section_title=f"{maintenance_index}. 维护与支持",
                index = maintenance_index,
                product_description=product_description,
                product_name=self.product_name,
                func_name=func_name,
                func_info=func_info_str
            )
            cleaned_paragraph = maintenance_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)
        except Exception as e:
            logger.error(f"Error generating maintenance and support section: {e}")
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")
            logger.error("Traceback details:", exc_info=True)
            

if __name__ == "__main__":
    # Define the list of platforms
    platforms = [
        {
            "docx_name": "多租户管理系统技术规范书.docx",
            "title": "多租户管理系统技术规范书",
            "product_name": "多租户管理系统"
        }
        
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
