import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.chains import LLMChain
import docx
from typing import Any
import logging
from langchain.chains import LLMChain
from docx.shared import Pt


# 大模型模块技术规范书生成器
# 一次性生成以下所有模块的技术规范书
# 模块: 
# 数据预处理平台， 
# 模型训练平台，  
# 模型保存与部署平台， 
# 模型评估平台， 
# 模型推理平台


# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的产品描述和平台名称，撰写针对指定模块的技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括系统架构、模块功能、输入输出、处理逻辑和部署要求。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            重要注意事项：

            专注于当前模块，不涉及其他模块功能。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            在描述技术内容时，提供尽可能具体的细节，例如关键组件、数据流和测试标准。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            
            本次生成是以段落为单位，根据user_prompt去生成对应的段落，而不是整个文档，同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{module_name}的技术规范书引言，简要介绍文档的目的和范围，不少于500字。
                    生成的不是技术规范书的整体, 而是针对该模块的引言段落。
                    该模块是以下产品的一部分, 请专注于该模块本身：{product_description}
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    示例:
                    1.1 目的
                    本文档的编制目的是为了定义智能客服系统的技术标准和规范，确保系统开发、部署及运维过程的规范性和一致性。通过对系统的功能、性能、接口、设计和实施要求的详细规定，旨在保障系统能够满足业务需求，同时提升系统的稳定性、安全性和用户体验。本文档将作为项目开发团队、测试团队、运维团队以及相关利益相关者的参考依据，确保各方对系统的理解和期望达成一致。
                    1.2 范围
                    本技术规范书全面覆盖智能客服系统的构建和运行所需的全部技术要求。具体包括但不限于：
                    - 系统架构设计，包括系统的总体设计思路、模块划分和数据流向。
                    - 功能实现细节，涉及系统的核心功能点以及附加功能的具体实现方式。
                    - 接口标准规定，明确系统与外部系统集成时的接口规范和数据交互格式。
                    - 安全性要求，确保系统在数据存储、传输和处理过程中的安全性。
                    - 可靠性指标，保障系统的高可用性、故障恢复能力和性能稳定性。
                    - 维护和支持指南，提供系统的维护流程、升级策略和技术支持渠道。
                    1.3 参考文档
                    为确保本技术规范书的科学性和权威性，编制过程中参考了以下文档和标准：
                    - ISO/IEC 25010:2011 系统和软件工程——系统和软件产品质量模型，提供了系统和软件质量的通用评价框架。
                    - 相关API接口文档，为系统提供了与外部服务和系统集成所需的接口信息和数据交互标准。
                    - 数据保护法规，包括GDPR、CCPA等，确保系统在处理用户数据时遵守相应的法律法规。
                    1.4 定义和缩略语
                    为了确保本文档中的术语使用一致和清晰，以下列出了一些关键词汇及其定义：
                    - CRM（Customer Relationship Management）：客户关系管理系统，用于帮助企业管理与客户的互动信息。
                    - API（Application Programming Interface）：应用程序编程接口，允许不同软件组件之间进行互相通信。
                    - SDK（Software Development Kit）：软件开发工具包，提供了一套工具、指南和API，以便开发者创建软件应用程序。
                    - UI（User Interface）：用户界面，指系统与用户交互的界面部分，包括布局、视觉元素和交互逻辑。
                    - UX（User Experience）：用户体验，指用户在使用系统过程中的总体感受和体验。
                    - HTTPS（Hypertext Transfer Protocol Secure）：安全超文本传输协议，通过对传输数据进行加密，提供了数据传输的安全性保障。
                    - SSL/TLS（Secure Sockets Layer / Transport Layer Security）：安全套接层/传输层安全，两种用于在互联网上确保数据传输安全的协议。
                    - NLP（Natural Language Processing）：自然语言处理，指使计算机能够理解和处理人类语言的技术。

                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{module_name}的系统概述，包括其主要功能和用途, 不少于600字。
                    生成的不是技术规范书的整体, 而是针对该模块的系统概述段落，是整个技术规范书的第二段，第一段是引言，所以专注于系统概述本身，不需要涉及引言或结论等段落。
                    该模块是以下产品的一部分, 请专注于该模块本身：{product_description}
                    示例:
                    2.1 系统功能
                    智能客服系统旨在整合多种通信渠道和客服资源，提供高效、便捷的客户服务体验。以下是系统的主要功能：
                    - 多渠道接入：
                    - 桌面网站：通过提供可嵌入网页的聊天插件，使用户能够直接在网站上与客服交流。
                    - 移动网站：确保聊天插件能够在移动设备上自动适配屏幕大小，保持良好的用户体验。
                    - App：提供适用于iOS和Android平台的SDK，使得App内可以无缝接入客服功能。
                    - 微信：集成微信公众号和小程序的客服接口，直接在微信环境中提供客服服务。
                    - 微博：通过微博开放平台实现客服功能，与用户进行互动。
                    - 短信：接入短信服务平台，为无法使用即时通信的用户提供客服支持。
                    - 接待方式：
                    - 人工客服：系统为人工客服提供完整的操作界面，包括即时消息处理、用户信息查看、历史记录访问等。
                    - 机器人客服：集成自然语言处理技术，实现对用户咨询的自动响应，并提供转接人工客服的选项。
                    - 技能组管理：管理员可以创建不同的技能组，根据客服人员的专长将其分配至相应组别，系统根据用户咨询内容自动分配至对应技能组。
                    - 富媒体沟通：
                    - 文字：支持发送接收即时文本消息。
                    - 表情：集成表情包，丰富交流情感表达。
                    - 图片：允许发送图片文件，并在聊天窗口中直接预览。
                    - 富文本：支持发送包含格式的文本消息，如加粗、斜体、列表等。
                    - 超链接：允许在消息中插入网页链接，用户点击后可直接跳转。
                    - 会话管理：
                    - 自动应答：在客服不在线或忙碌时，系统自动回复预设的消息，引导用户或提供基本帮助。
                    - 自动弹框邀请：根据用户在网站的行为设置自动邀请策略，主动提供帮助。
                    - 主动邀请：客服可以根据用户行为主动发起聊天邀请，提高互动率。
                    - 聊天信息同步：系统保留用户与不同客服的聊天记录，确保信息的连续性和完整性。
                    - 用户分析：
                    - 用户画像：展示用户的基本信息，如昵称、联系方式、历史订单等，帮助客服提供个性化服务。
                    - 访问轨迹：记录用户在网站的访问路径，分析用户兴趣和需求。
                    - 用户访问分析：提供用户访问数据的报表，支持自定义查询和导出。
                    - 会话记录：允许管理员查询历史会话记录，用于服务质量评估和监控。
                    - 客户满意度：
                    - 评价系统：用户可对服务进行满意度评价，收集反馈用于改进服务。
                    2.2 用户特征
                    智能客服系统面向以下用户群体：
                    - 客服人员：使用系统接待访客，处理咨询和问题，需要系统提供稳定、高效的操作界面。
                    - 系统管理员：负责系统的设置、维护和升级，需要管理工具来监控系统状态和性能。
                    - 访客：通过不同渠道接触系统，寻求即时的帮助和信息，期待快速、准确的服务。
                    2.3 运行环境
                    智能客服系统的运行环境要求如下：
                    - 服务器端：
                    - 操作系统：支持Linux和Windows服务器操作系统。
                    - 部署：支持物理服务器、虚拟化环境及云服务平台部署。
                    - 数据库：兼容主流数据库系统，如MySQL、PostgreSQL等。
                    - 客户端：
                    - 浏览器：确保与主流浏览器如Chrome、Firefox、Safari和Edge兼容。
                    - 移动操作系统：支持iOS和Android操作系统，确保App内客服功能的稳定运行。

                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{module_name}的系统架构, 不少于800字。
                    生成的不是技术规范书的整体, 而是针对该模块的系统架构段落。
                    该模块是以下产品的一部分, 请专注于该模块本身：{product_description}
                    架构描述应包括数据流、处理步骤和关键组件。
                    示例:
                    3.1 总体架构
                    智能客服系统将采用一种模块化、分布式的架构设计，以确保系统的可扩展性、可维护性和高可用性。系统的主要架构分为以下几个层次：
                    3.1.1 前端展示层
                    前端展示层负责提供用户界面(UI)，使用户能够通过图形界面与系统进行交互。该层主要包括：
                    - 客户端UI：为不同终端（桌面、移动、App等）提供定制化的用户界面。
                    - Web UI：为桌面和移动浏览器提供响应式设计的网页界面。
                    - 第三方平台UI：为微信、微博等第三方平台集成的客服界面。
                    3.1.2 业务逻辑层
                    业务逻辑层是系统的核心，负责处理客服业务流程和用户交互逻辑。包括但不限于：
                    - 会话管理：处理用户与客服的会话，包括会话的创建、维持、结束和转接。
                    - 技能组路由：根据用户的问题和预设的规则，将用户的咨询定向到合适的技能组。
                    - 消息处理：处理发送和接收的消息，包括文本、图片、视频等多种格式。
                    
                    3.2 系统模块划分
                    系统根据功能需求被划分为多个模块，以实现职责的分离和功能的独立。
                    3.2.1 用户管理模块
                    - 用户认证：负责用户登录认证过程，包括密码校验、会话管理等。
                    - 用户权限：管理用户的权限设置，确保用户只能访问授权的资源。
                    - 用户信息：维护用户的基本信息，包括联系方式、偏好设置等。
                    3.2.2 消息处理模块
                    - 消息队列：采用消息队列进行消息的收发，保证消息的可靠传输。
                    - 消息格式化：将消息转换为标准格式，以便在不同的通信渠道中传输。
                    - 消息记录：记录所有的消息历史，以供查询和分析。
                    3.2.3 数据分析模块
                    - 访问分析：分析用户的访问行为和路径，生成访问报告。
                    - 会话分析：统计会话的相关数据，如持续时间、转化率等。
                    - 用户满意度：收集和分析用户满意度调查的结果。
                    
                    """
                )
            ),
            
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name", "func_name"],
                template=(
                    """
                    {system_message}
                    请详细描述功能“{func_name}”的技术规范段落，不少于600字。
                    
                    该功能是模块“{module_name}”的重要组成部分，所以生成的不是技术规范书的整体，而是针对该功能的技术规范段落。
                    模块“{module_name}”又是以下产品的一部分：{product_description}。
                    请严格聚焦于该功能，避免涉及其他模块或功能。

                    技术规范应包括以下内容：
                    1. **功能描述**：概述功能的核心目标、输入输出及其处理流程。
                    2. **技术要求**：列出该功能的关键技术需求，例如性能指标、可用性、响应时间等。
                    3. **依赖关系**：明确该功能是否依赖其他系统、模块或资源。
                    4. **实现方案**：提供高层次的实现细节，包括使用的技术框架、算法或工具, 具体到版本。
                    5. **部署要求**：说明功能的运行环境需求，如硬件、操作系统、依赖库等。
                    6. **扩展性与可维护性**：描述功能在未来的扩展和维护方面的设计考虑。

                    在撰写时，请以清晰、结构化的方式逐点展开，确保内容详尽且专业。
                    
                    以下是一个示例:
                    4. 接入渠道技术规范
                    4.1 桌面网站接入
                    4.1.1 JavaScript SDK
                    - 提供JavaScript SDK以供集成到企业桌面网站中，该SDK应包含以下功能：
                    - 用户身份验证：确保用户与系统之间的通信是安全的。
                    - 会话管理：允许创建、维护和结束用户会话。
                    - 消息发送和接收：支持文本、图片、文件等多种消息类型的发送和接收。
                    - 事件监听：能够响应用户交互事件，如点击、输入等。
                    - UI组件：提供标准的聊天窗口UI组件，支持自定义样式以符合企业品牌形象。
                    4.1.2 通讯机制
                    - 系统应支持WebSocket协议，以实现实时、双向的通信。
                    - 对于不支持WebSocket的环境，系统应提供长轮询的备选方案，以保证消息的可靠传输。
                    - 所有通讯应通过HTTPS进行加密，以确保数据传输的安全性。
                    4.2 移动网站接入
                    4.2.1 自适应设计
                    - 移动端接入应使用响应式设计，确保UI组件能够根据不同屏幕尺寸进行适配。
                    - SDK应提供与桌面网站接入相同的功能，并优化触摸操作和网络条件变化的响应。
                    4.3 App接入
                    4.3.1 移动SDK
                    - 提供适用于iOS和Android平台的SDK，该SDK应包括以下功能：
                    - 原生UI组件：提供原生操作系统风格的聊天窗口组件。
                    - 推送通知：支持系统消息推送，包括新消息通知和会话更新通知。
                    - 网络管理：智能处理网络变化，如断网、重连等情况。
                    4.3.2 接口文档
                    - 提供详尽的SDK接口文档，包括安装指南、功能说明、示例代码和常见问题解答。
                    4.4 微信接入
                    4.4.1 公众号和小程序
                    - 集成微信公众号和小程序的客服功能，实现以下接口：
                    - 消息接收：接收用户在微信中发送的消息。
                    - 消息回复：允许客服回复用户消息，并支持消息类型多样化。
                    - 用户事件处理：处理用户的关注、取消关注等事件。
                    4.4.2 微信API集成
                    - 使用微信官方提供的API完成系统集成，确保功能的稳定性和兼容性。
                    4.5 微博接入
                    4.5.1 微博客服接口
                    - 通过微博开放平台提供的API实现客服功能，包括：
                    - 用户验证：验证用户身份，确保消息来源的真实性。
                    - 消息交互：实现在微博平台上的消息发送和接收。
                    4.5.2 接口兼容性
                    - 确保与微博平台的接口兼容性，及时更新以适应微博API的变化。
                    4.6 短信接入
                    4.6.1 短信服务商合作
                    - 选择合规的第三方短信服务商，提供短信发送和接收服务。
                    - 确保短信服务的覆盖范围广泛，支持主要运营商和国际短信。
                    4.6.2 短信接口集成
                    - 集成短信发送和接收的API，实现以下功能：
                    - 短信发送：允许系统通过API发送短信给用户。
                    - 短信接收：处理用户回复的短信，并接入系统进行相应处理。
                    - 短信格式：支持自定义短信内容，包括签名和模板。
                    """
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "module_name"],
                template=(
                    """
                    {system_message}
                    请描述{module_name}的维护和支持策略，包括维护计划和支持渠道,不少于400字。
                    维护策略应说明{module_name}的维护计划和流程。
                    生成的不是技术规范书的整体, 而是针对该模块的维护与支持段落。
                    该模块是以下产品的一部分, 请专注于该模块本身：{product_description}
                    """
                )
            )
            
            
        }

    def generate_text(self, prompt_template, product_description, module_name, func_name = ""):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == "":
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                module_name=module_name,
            )
        else:
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                module_name=module_name,
                func_name=func_name
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text


    def generate_section(self, prompt_template, section_title, product_description, module_name, func_name = "", **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            product_description=product_description,
            module_name=module_name,
            func_name=func_name
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()

        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction"]},
        {"title": "系统概述", "prompt": self.prompts["system_overview"]},
        {"title": "系统架构", "prompt": self.prompts["system_architecture"]},
        ]

        # Step 4: 生成前三个部分
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                product_description=product_description,
                module_name=self.product_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('*', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

        # Step 5: 提取功能模块名称
        module_names_prompt = (
            f"根据以下产品简述，提取出适合的功能模块名称列表，每个模块名称应准确描述其功能，列表中应指包括模块名称，不要有多余开头或结尾"
            f"适用于技术规范书的编写，结果以列表形式返回：\n{product_description}"
        )
        module_names_response = self.llm.predict(module_names_prompt)
        module_names = module_names_response.strip().split('\n')
        print(f"功能模块有: {module_names}")

        # Step 6: 为每个功能模块生成技术规范
        for index, func_name in enumerate(module_names, start=4):  # 从第 4 章开始
            logger.info(f"生成功能技术规范: {func_name}")
            try:
                section_paragraph = self.generate_section(
                    prompt_template=self.prompts["technical_specification"],
                    section_title=f"{index}. {func_name}技术规范",
                    product_description=product_description,
                    module_name=self.product_name,
                    func_name=func_name
                )
                # 清理多余字符并写入文档
                cleaned_paragraph = section_paragraph.replace('*', '').replace('#', '')
                document.add_paragraph(cleaned_paragraph)
            except Exception as e:
                logger.error(f"Error generating technical specification for {module_name}: {e}")

        
        # 添加维护与支持部分
        try:
            maintenance_index = len(module_names) + 4  # 最后一部分的索引
            maintenance_paragraph = self.generate_section(
                prompt_template=self.prompts["maintenance_support"],
                section_title=f"{maintenance_index}. 维护与支持",
                product_description=product_description,
                module_name=self.product_name
            )
            cleaned_paragraph = maintenance_paragraph.replace('*', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)
        except Exception as e:
            logger.error(f"Error generating maintenance and support section: {e}")
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")

if __name__ == "__main__":
    # Define the list of platforms
    product_name = "MoE算法V1.0"
    platforms = [
        {
            "docx_name": "MoE专家网络设计规范书.docx",
            "title": f"{product_name} - 专家网络设计规范书",
            "product_name": "专家网络设计"
        }
        # {
        #     "docx_name": "MoE门控机制规范书.docx",
        #     "title": f"{product_name} - 门控机制规范书",
        #     "product_name": "门控机制"
        # },
        # {
        #     "docx_name": "MoE训练策略规范书.docx",
        #     "title": f"{product_name} - 训练策略规范书",
        #     "product_name": "训练策略"
        # },
        # {
        #     "docx_name": "MoE模型评估与优化规范书.docx",
        #     "title": f"{product_name} - 模型评估与优化规范书",
        #     "product_name": "模型评估与优化"
        # },
        # {
        #     "docx_name": "MoE部署与应用规范书.docx",
        #     "title": f"{product_name} - 部署与应用规范书",
        #     "product_name": "部署与应用"
        # }
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
