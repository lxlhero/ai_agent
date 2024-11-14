import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import docx
from typing import Any
import logging
from langchain.chains import LLMChain


# Add after imports, before class definition
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ["OPENAI_API_KEY"] = "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"

class DocumentAgent:
    def __init__(self, overview_path, docx_name, title):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        
        # 系统prompt
        self.system_message = """
            你是一个文档编写者，你的任务是根据产品概述撰写技术规范书的段落, 
            技术规范书的要点是明确目的和范围、结构化和组织良好、详细和具体、准确性和可靠性，
            注意每个段落不要在开头出现段落title,比如撰写引言时，不要出现引言二字
            """
        
        # 引言
        self.introduction_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请编写一段关于大模型的技术规范书引言，简要介绍文档的目的和范围，不少于600字。
                请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                以下是产品简述：\n{product_description}
                """
            )
        )
        
        # 大模型概述
        self.large_model_overview_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请提供一个技术规范书的大模型概述，包括多模态、微调方式、用户特征和运行环境。功能概览应简要介绍大模型的主要功能和用途, 不少于800字。
                产品简述：{product_description}
                不少于800字。
                """
            )
        )
        
        # 数据集描述
        self.dataset_description_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请详细描述用于训练、验证和测试大模型的数据集，不少于500字。
                数据集描述应包括以下内容：
                1. 数据集的来源和采集方法。
                2. 数据集的组成和结构。
                3. 数据预处理和清洗步骤。
                4. 数据集的标注方法和质量评估。
                以下是产品简述：{product_description}
                """
            )
        )
        
        # 训练过程
        self.training_process_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请详细描述大模型的训练过程，不少于800字。
                训练过程描述应包括以下内容：
                1. 使用的硬件和软件环境。
                2. 训练数据的划分和使用。
                3. 训练参数的设置（如学习率、批次大小等）。
                4. 优化算法和正则化技术。
                5. 训练过程中的监控和评估指标。
                以下是产品简述：{product_description}
                """
            )
        )
        
        # 模型评估
        self.model_evaluation_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请详细描述大模型的评估方法和评估结果，不少于600字。
                模型评估描述应包括以下内容：
                1. 使用的评估指标和基准测试。
                2. 评估数据集的选择和使用。
                3. 评估结果的详细报告和分析。
                4. 模型的局限性和潜在改进方向。
                以下是产品简述：{product_description}
                """
            )
        )
        
        # 部署指南
        self.deployment_guide_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请提供大模型的部署指南，不少于600字。
                部署指南应包括以下内容：
                1. 部署所需的硬件和软件要求。
                2. 部署步骤和配置说明。
                3. 性能调优和资源管理建议。
                4. 常见问题和解决方案。
                以下是产品简述：{product_description}
                """
            )
        )
        
        #风险管理
        self.risk_management_prompt = PromptTemplate(
            input_variables=["productация_description"],
            template=(
                """
                请识别和评估与大模型相关的潜在风险，并提出应对措施，不少于500字。
                风险管理描述应包括以下内容：
                1. 数据隐私和安全风险。
                2. 模型偏见和不公平性。
                3. 系统故障和可靠性问题。
                4. 法律和合规风险。
                以下是产品简述：{product_description}
                """
            )
        )
        
        # 维护和支持
        self.large_model_maintenance_support_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """请描述大模型的维护和支持策略，包括维护计划和支持渠道。维护策略应说明大模型的维护计划和流程。支持渠道应提供用户和技术支持的联系信息和渠道。
                产品简述：{product_description}。
                不少于500字
                """
            )
        )
        
        # 未来发展和迭代
        self.large_model_future_development_iteration_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请描述大模型未来的技术发展方向和改进计划，以及大模型的扩展能力和策略。
                技术迭代应说明大模型未来的技术发展方向和改进计划。
                扩展性应描述大模型的扩展能力和策略。
                产品简述：{product_description}。
                不少于500字
                """
            )
        )
        
        # 结论
        self.large_model_conclusion_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请编写结论部分，简要总结文档的主要内容和要点。
                总结应概述文档的主要内容和要点。
                下一步行动应说明文档发布后的下一步行动计划。
                产品简述：{product_description}。
                不少于500字
                
                """
            )
        )
        
        # 附录
        self.large_model_appendix_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                """
                请编写附录部分，包括术语表和参考资料。
                术语表应提供文档中使用的专业术语的解释。参考资料应列出相关的标准、法规和其他参考文献。
                """
            )
        )

    
    def generate_text(self, prompt_template, **kwargs):
        """
            Generate text using LLMChain with the provided prompt template.
        """
        chain = LLMChain(prompt=prompt_template, llm=self.llm)
        generated_text = chain.run(**kwargs)
        
        logger.info("Generated Text: %s", generated_text)
        return generated_text


    def generate_section(self, document, prompt_template, section_title, product_description,  **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        kwargs['system_message'] = self.system_message
        
        # Generate text content for the section
        section_content = self.generate_text(prompt_template, product_description=product_description,)
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n\n{section_content}"
    
        return section_string
    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()

        

        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
            {"title": "引言", "prompt": self.introduction_prompt},
            {"title": "大模型概述", "prompt": self.large_model_overview_prompt},
            {"title": "数据集描述", "prompt": self.dataset_description_prompt},
            {"title": "训练过程", "prompt": self.training_process_prompt},
            {"title": "模型评估", "prompt": self.model_evaluation_prompt},
            {"title": "部署指南", "prompt": self.deployment_guide_prompt},
            {"title": "风险管理", "prompt": self.risk_management_prompt},
            {"title": "维护和支持", "prompt": self.large_model_maintenance_support_prompt},
            {"title": "未来发展和迭代", "prompt": self.large_model_future_development_iteration_prompt},
            {"title": "结论", "prompt": self.large_model_conclusion_prompt},
            {"title": "附录", "prompt": self.large_model_appendix_prompt}
        ]

        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(document, section["prompt"], f"{index}. {section['title']}", product_description)
            document.add_paragraph(section_paragraph)

        # 保存文档
        document.save(self.docx_name)
        logger.info(f"技术规范书已保存为 {self.docx_name}")

if __name__ == "__main__":
    agent = DocumentAgent(
        "./overview.txt",
        "康养行业模型技术规范书.docx",
        "康养行业模型技术规范书"
    )
    agent.process()