import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import docx
from typing import Any
import logging
from langchain.chains import LLMChain


# Add after imports, before class definition
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ["OPENAI_API_KEY"] = "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"

class DocumentAgent:
    def __init__(self, overview_path, modules_path, docx_name, title):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.modules_path = modules_path
        self.docx_name = docx_name
        self.title = title
        
        # 系统prompt
        self.system_message = """
            你是一个文档编写者，你的任务是根据产品概述撰写技术规范书的段落, 
            技术规范书的要点是明确目的和范围、结构化和组织良好、详细和具体、准确性和可靠性，
            注意每个段落不要在开头出现段落title,比如撰写引言时，不要出现引言二字
            """
        
        # 引言
        self.introduction_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "请编写一段技术规范书引言，简要介绍文档的目的和范围。请包括文档的目标、涵盖的范围以及相关的参考文档和定义。以下是产品简述：\n{product_description}"
            )
        )
        # 系统概述
        self.system_overview_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """
                请提供一个技术规范书的系统概述，包括功能概览、用户特征和运行环境。功能概览应简要介绍系统的主要功能和模块。
                产品简述：{product_description}
                \n用户特征应描述不同用户角色（如管理员、用户、开发者）的需求和特征。运行环境应说明系统所需的硬件和软件环境。
                以下是模块功能简介,请根据此信息完善系统概述: {modules_info}
                """
            )
        )
        
        # 系统架构
        self.system_architecture_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """
                请详细描述系统的总体架构和各个模块的设计。
                总体架构应包括数据层、服务层、应用层和表示层的职责。模块划分应详细介绍各个功能模块的设计和职责。
                以下是模块功能简介,请根据此信息完善系统架构: {modules_info}
                """
            )
        )
        
        # 功能模块详细说明
        self.function_module_detail_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """请详细说明系统的各个功能模块, 包括接入渠道、核心功能和智能功能。接入渠道应描述系统支持的接入方式（如网页、移动应用、社交媒体）。
                核心功能应介绍系统的基本功能模块（如用户管理、数据处理、业务逻辑）。智能功能应介绍系统的智能化功能（如AI、机器学习、数据分析）。
                以下是模块功能简介,请根据此信息完善功能模块详细说明:{modules_info}
                """
            )
        )
        
        # 技术规范
        self.technical_specification_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """
                请描述系统的技术规范，包括数据管理、安全性和性能要求。
                数据管理应说明数据的存储、备份和恢复策略。安全性应描述系统的安全措施（如加密、访问控制）。
                性能要求应定义系统的性能指标（如响应时间、并发用户数）。
                以下是模块功能简介,请根据此信息完善技术规范{modules_info}
                """
            )
        )   
        
        # 用户界面和体验
        self.user_interface_experience_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """
                请描述系统的用户界面设计和用户体验。界面设计应说明用户界面的设计原则和交互方式。
                用户体验应说明如何提升用户的使用体验（如响应速度、界面友好性）。
                以下是模块功能简介,请根据此信息完善用户界面和体验: {modules_info}
                """
            )
        )
        
        # 测试和验证
        self.testing_validation_prompt = PromptTemplate(
            input_variables=["product_description", "modules_info"],
            template=(
                """
                请编写测试和验证部分的文档，包括测试计划、验证和验收标准。测试计划应描述系统的测试策略和测试用例。
                验证和验收标准应说明系统的验证和验收标准。
                以下是模块功能简介,请根据此信息完善测试和验证: {modules_info}
                """
            )
        )
        
        # 维护和支持
        self.maintenance_support_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "请描述系统的维护和支持策略，包括维护计划和支持渠道。维护策略应说明系统的维护计划和流程。支持渠道应提供用户和技术支持的联系信息和渠道。"
            )
        )
        
        # 未来发展和迭代
        self.future_development_iteration_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "请描述系统未来的技术发展方向和改进计划，以及系统的扩展能力和策略。技术迭代应说明系统未来的技术发展方向和改进计划。扩展性应描述系统的扩展能力和策略。"
            )
        )
        
        # 结论
        self.conclusion_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "请编写结论部分，简要总结文档的主要内容和要点。总结应概述文档的主要内容和要点。下一步行动应说明文档发布后的下一步行动计划。"
            )
        )
        
        # 附录
        self.appendix_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "请编写附录部分，包括术语表和参考资料。术语表应提供文档中使用的专业术语的解释。参考资料应列出相关的标准、法规和其他参考文献。"
            )
        )

    
    def generate_text(self, prompt_template, **kwargs):
        """
            Generate text using LLMChain with the provided prompt template.
        """
        chain = LLMChain(prompt=prompt_template, llm=self.llm)
        generated_text = chain.run(**kwargs)
        
        logger.info("Generated Text: %s", generated_text)
        return generated_text


    def generate_section(self, document, prompt_template, section_title, product_description, modules_info, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        prompt = prompt_template.format(product_description=product_description, modules_info=modules_info)
        kwargs['system_message'] = self.system_message
        
        # Generate text content for the section
        section_content = self.generate_text(prompt_template, product_description=product_description, modules_info=modules_info)
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n\n{section_content}"
    
        return section_string
    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()

        # 读取并整理产品的功能模块
        with open(self.modules_path, 'r', encoding='utf-8') as file:
            modules = json.load(file)

        modules_list = []

        # 遍历字典并将每个模块的信息格式化为字符串，然后添加到列表中
        for index, (module_number, module_info) in enumerate(modules.items(), start=1):
            module_string = f"{index}. 功能: {module_info['模块名称']}\n   说明: {module_info['基本功能说明']}\n"
            modules_list.append(module_string)

        # 产品的功能模块信息
        modules_info = '\n'.join(modules_list)

        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
            {"title": "引言", "prompt": self.introduction_prompt},
            {"title": "系统概述", "prompt": self.system_overview_prompt},
            {"title": "系统架构", "prompt": self.system_architecture_prompt},
            {"title": "功能模块详细说明", "prompt": self.function_module_detail_prompt},
            {"title": "技术规范", "prompt": self.technical_specification_prompt},
            {"title": "用户界面和体验", "prompt": self.user_interface_experience_prompt},
            {"title": "测试和验证", "prompt": self.testing_validation_prompt},
            {"title": "维护和支持", "prompt": self.maintenance_support_prompt},
            {"title": "未来发展和迭代", "prompt": self.future_development_iteration_prompt},
            {"title": "结论", "prompt": self.conclusion_prompt},
            {"title": "附录", "prompt": self.appendix_prompt}
        ]

        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(document, section["prompt"], f"{index}. {section['title']}", product_description, modules_info)
            document.add_paragraph(section_paragraph)

        # 保存文档
        document.save(self.docx_name)
        logger.info(f"技术规范书已保存为 {self.docx_name}")

if __name__ == "__main__":
    agent = DocumentAgent(
        "./智能客服系统概览.txt",
        "./智能客服系统功能需求.txt",
        "智能客服系统技术规范书.docx",
        "智能客服系统技术规范书"
    )
    agent.process()