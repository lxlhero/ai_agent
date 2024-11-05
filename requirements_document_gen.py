# 根据产品简述和产品需求分解excel,构建agent生成需求文档

import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import bs4
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import docx
from typing import Any





os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = "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"
os.environ["OPENAI_API_KEY"] = "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"

class DocumentAgent:
    def __init__(self, overview_path, excel_path, docx_name, title):
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        # 产品简述文件地址
        self.overview_path = overview_path
        # 产品功能excel地址
        self.excel_path = excel_path
        # 需求文档名称
        self.docx_name = docx_name
        # 需求文档标题
        self.title = title
        
        
        # 项目背景和产品需求综述
        self.summary_prompt = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "以下是一个产品的简述: {product_description}\n"
                "请基于这个简述生成产品需求文档的开头综述。"
                "概述该产品的背景与目标, 产品愿景, 适用范围, 不少于500字"
            )
        )
        
     

    # Execute the chain to get the product summary
    def generate_product_summary(self, product_description):
        """
        生成产品摘要。
        
        Args:
            product_description (str): 产品描述。
        
        Returns:
            str: 生成的产品摘要。
        
        """
        # Run the chain with the product description to get a summary
        product_summary = self.summary_chain.run(product_description=product_description)
        return product_summary
    
    
    def search_module(self, chroma_db, module_name):
        """
        在指定的色谱数据库中搜索模块。
        
        Args:
            chroma_db (ChromaDatabase): 色谱数据库对象，用于执行搜索查询。
            module_name (str): 要搜索的模块名称。
        
        Returns:
            list: 包含最多1个搜索结果的列表。
        
        """
        # Perform the search query
        results = chroma_db.similarity_search(
            query=module_name,  # Search for the module name
            k=1  # Specify the number of results to return
        )
        
        return results
    
    # Function to format the product description
    def format_product_description(self, product_description):
        return f"产品简述: {product_description}"


    def process(self):

        llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述, 作为之后prompt的上下文
        with open(self.overview_path, 'r', encoding='utf-8') as f:
            product_description = f.read()
            
        # 初始化LLMChain，用于生成产品需求文档开头的综述
        self.summary_chain = LLMChain(
        llm=llm,  # Use the LLM initialized in your process
        prompt=self.summary_prompt
        )
        
                    
        product_summary = self.generate_product_summary(product_description)
        print("Product Summary:", product_summary)
        
        # 创建一个新的Word文档
        doc = docx.Document()
        # 添加标题（可根据需要修改）
        doc.add_heading(self.title, level=1)
        # 将产品总结添加到文档中
        doc.add_paragraph(product_summary)
        # 保存文档，这里假设保存路径为'summary.docx'，你可以根据需求修改
        doc.save(self.docx_name)
        
        

        # 从excel中读取产品功能，存入本地知识库，为后续RAG搭建做准备
        
        # sheets = pd.read_excel(self.excel_path, sheet_name=None)
        # df = sheets[next(iter(sheets))]

        # # 模块list，待会根据这个从db语义查询出相关数据做context
        # self.modules = df['模块'].unique().tolist()

        # # 以模块为维度插入db
        # grouped = df.groupby('模块')['功能'].apply(lambda x: ', '.join(x)).reset_index()

        # # Prepare and insert into Chroma DB
        # for _, row in grouped.iterrows():
        #     module = row['模块']
        #     functions = row['功能']
        #     # Make the content more descriptive
        #     doc_content = f"这是一个关于{module}模块的文档，其中功能包括: {functions}。"
        #     self.chroma_db.add_texts([doc_content])
            
        


        
            
            
            
            
if __name__ == "__main__":
    agent = DocumentAgent("./description.txt", "./客服会话智能分析预警系统.xlsx", "客服会话智能分析预警系统需求文档.docx", "客服会话智能分析预警系统需求文档")
    agent.process()


    

    





    
