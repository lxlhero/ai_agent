import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
import docx
from typing import Any
import logging
from docx.shared import Pt
import traceback
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment



# 大模型模块技术规范书生成器
# 根据产品简述生成技术规范书



# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        # 这里是整个产品的名称
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的数据库服务技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的数据库简述，撰写技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括系统架构、模块功能、输入输出、处理逻辑和部署要求。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            重要注意事项：
            规范书的内容主要是数据库软件的一些说明内容。
            user_prompt中的数据库简述提供该产品的背景知识和功能，帮助你去撰写技术规范书。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            在描述技术内容时，提供尽可能具体的细节。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            本次生成是以段落为单位,根据user_prompt去生成对应的段落,而不是整个文档,同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可。
            注意每个段落的目录分级，例如第一个段落引言，需要列出1.1, 1.2等目录，每个段落都需要目录。
            
            应该使用规定性语言如应该、不得、提供、应支持等，以增强文档的权威性和规范性。
            应该多使用英文术语以提高规范性和权威性。
            语言风格示例:
            创建技能组
             提供一个用户友好的界面，允许管理员创建新的技能组。
             管理员可以为每个技能组指定一个唯一的名称和描述，以便于识别和管理。
             技能组的创建过程应支持自定义配置，如服务时间、语言能力、专业领域等。
            管理技能组
             管理员可以编辑已存在的技能组信息，包括修改名称、描述和配置参数。
             提供技能组的启用和禁用功能，允许临时停用某个技能组而不删除其配置。
             技能组配置应支持版本控制，记录历史更改并允许恢复到之前的配置。
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{product_name}的技术规范书引言，简要介绍文档的目的和范围，不少于1000字。
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    生成的不是技术规范书的整体, 而是针对该模块的引言段落。
                    
                    以下是产品简述：{product_description}
                    
                    以下是简短的示例:
                     1.1 目的
                    本文档旨在明确数据库技术规范，确保数据库的设计、实施、运行和维护过程遵循统一标准。通过详细规定数据库的性能、安全性、可用性和扩展性要求，旨在提供高效、稳定及安全的数据管理能力，以支持业务需求和提升用户体验。本文档作为开发、运维及相关团队的指导性文件，确保项目各方对数据库的技术要求和实现细节有共同的理解。

                        1.2 范围
                    本技术规范书涵盖数据库选型、架构设计、性能优化、安全策略、备份恢复、监控及维护等全方位要求。具体包括数据库类型的选取、系统架构的设计思路、数据存储与访问的性能指标、数据安全和隐私保护措施、高可用性和容灾能力的实现方法，以及日常运维和技术支持的具体指南。

                        1.3 参考文档
                    编写本技术规范书时，参考了以下标准和文献：
                    - ISO/IEC 25010:2011 系统和软件工程——系统和软件产品质量模型。
                    - 相关数据库管理和优化最佳实践文档。
                    - 国家和地区的数据安全和隐私保护法律法规。

                        1.4 定义和缩略语
                    为确保文档术语的准确性和一致性，以下是部分关键术语及其定义：
                    - DBMS（Database Management System）：数据库管理系统，用于创建、维护和管理数据库。
                    - ACID（Atomicity, Consistency, Isolation, Durability）：描述事务处理特性的四个条件。
                    - RAID（Redundant Array of Independent Disks）：独立磁盘冗余阵列，一种提升数据可靠性和性能的存储技术。
                    - ORM（Object-Relational Mapping）：对象关系映射，一种将对象模型映射到关系数据库的技术。
                    - SQL（Structured Query Language）：结构化查询语言，用于管理关系数据库的标准编程语言。

                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{product_name}的系统概述，包括其主要功能和用途, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统概述段落，是整个技术规范书的第二段，第一段是引言，所以专注于系统概述本身，不需要涉及引言或系统架构等段落。
                    以下是产品简述：{product_description}
                    
                    以下是简短的示例:
                     2.1 数据库系统架构概述
                    数据库系统旨在提供高效、稳定且安全的数据存储与访问服务。以下是系统的主要架构特点：

                    2.1.1 数据库类型选择
                    关系型数据库：如MySQL、PostgreSQL，适用于结构化数据存储，支持复杂查询和事务处理。
                    非关系型数据库：如MongoDB、Redis，适用于非结构化或半结构化数据，提供高性能和灵活性。
                    向量数据库：如Milvus，专为机器学习和向量数据处理设计，优化了相似性搜索性能。

                    2.1.2 数据存储与管理
                    数据分区与分片：通过逻辑划分和物理分割，提高数据处理效率和系统扩展性。
                    数据复制与同步：实现数据的多个副本管理，确保数据的高可用性和容错性。
                    备份与恢复策略：定期执行数据备份，提供快速的数据恢复机制，保障数据安全。

                    2.1.3 数据访问与接口
                    API接口：提供标准化的应用程序接口，支持与其他系统的集成和数据交换。
                    查询语言支持：如SQL，提供强大的数据查询和分析能力。

                    2.1.4 数据安全与合规
                    访问控制：实施严格的用户认证和权限管理，保护数据不被未授权访问。
                    数据加密：在传输和存储过程中对敏感数据进行加密，遵守相关法律法规要求。

                    2.1.5 性能优化
                    缓存策略：利用内存数据库或缓存技术，减少数据库负载，提高响应速度。
                    索引优化：合理设计索引结构，提升查询效率。

                    2.2 用户与角色
                    数据库系统面向以下用户和角色：
                    数据库管理员：负责系统的配置、监控和维护，确保数据库的高性能和安全性。
                    开发人员：使用数据库进行应用开发，需要高效的开发工具和良好的数据库性能。
                    业务用户：通过应用程序间接访问数据库，获取所需数据和服务。

                    2.3 运行环境
                    数据库系统的运行环境要求如下：
                    服务器端：
                    操作系统：兼容主流的Linux发行版和Windows Server。
                    硬件：支持高性能服务器硬件配置，包括SSD存储和高速网络接口。
                    虚拟化与云服务：支持在虚拟化平台和云服务环境中部署，如AWS、Azure、阿里云等。
                    客户端：
                    开发工具：兼容主流的数据库管理工具和开发环境，如MySQL Workbench、pgAdmin等。
                    网络环境：支持广域网和局域网访问，确保数据的远程访问能力。

                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{product_name}的系统架构, 不少于1500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统架构段落。
                    架构描述应包括数据库类型和模型，数据库布局和结构, 数据分布和分片, 数据访问和接口等。
                    根据你对系统架构的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    以下是产品简述：{product_description}
                    
                    以下是简短的示例:
                     3.1 总体架构
                    数据库系统采用模块化、分布式的架构设计，以确保系统的可扩展性、可维护性和高可用性。系统的主要架构分为以下几个层次：

                    3.1.1 数据存储层
                    数据存储层负责数据的持久化存储和管理。该层主要包括：
                    关系型数据库：如MySQL、PostgreSQL，用于存储结构化数据，支持复杂查询和事务处理。
                    非关系型数据库：如MongoDB、Redis，用于存储非结构化或半结构化数据，提供高性能和灵活性。
                    向量数据库：如Milvus，专为机器学习和向量数据处理设计，优化了相似性搜索性能。

                    3.1.2 数据访问层
                    数据访问层提供数据的访问和管理接口，确保数据的有效检索和更新。该层主要包括：
                    数据访问对象（DAO）：提供对数据库的访问接口，封装数据操作逻辑。
                    ORM框架：如Hibernate、MyBatis，简化数据库操作，支持对象关系映射。
                    查询优化器：优化查询性能，确保高效的数据检索。

                    3.1.3 数据处理层
                    数据处理层负责数据的处理和分析，包括数据的转换、清洗和分析。该层主要包括：
                    ETL工具：如Apache NiFi、Talend，用于数据的抽取、转换和加载。
                    数据分析工具：如Apache Spark、Presto，用于大数据处理和分析。
                    数据仓库：如Amazon Redshift、Google BigQuery，用于存储和分析大规模数据集。

                    3.2 系统模块划分
                    数据库系统根据功能需求被划分为多个模块，以实现职责的分离和功能的独立。

                    3.2.1 用户管理模块
                    用户认证：负责用户登录认证过程，包括密码校验、会话管理等。
                    用户权限：管理用户的权限设置，确保用户只能访问授权的资源。
                    用户信息：维护用户的基本信息，包括联系方式、偏好设置等。

                    3.2.2 数据安全模块
                    访问控制：实施严格的用户认证和权限管理，保护数据不被未授权访问。
                    数据加密：在传输和存储过程中对敏感数据进行加密，遵守相关法律法规要求。
                    审计日志：记录所有数据库操作日志，以便追踪和审计。

                    3.2.3 数据备份与恢复模块
                    备份策略：定期执行数据备份，提供快速的数据恢复机制，保障数据安全。
                    恢复工具：如Percona XtraBackup、pg_dump，用于数据的恢复和迁移。
                    灾难恢复计划：制定详细的灾难恢复计划，确保在系统故障时能够快速恢复服务。

                    3.2.4 性能监控模块
                    监控工具：如Prometheus、Grafana，用于实时监控数据库性能和健康状况。
                    告警系统：设置告警规则，及时通知管理员处理异常情况。
                    性能优化：定期进行数据库性能调优，确保高效的数据处理能力。

                    
                    """
                )
            ),
            
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index", "func_name", "func_info"],
                template=(
                    """
                    {system_message}
                    请详细描述数据库“{func_name}”的技术规范段落, 不少于1800字。
                    这是该数据库的具体子功能和描述，请参考：“{func_info}”
                    技术规范需要涵盖该数据库的所有子功能技术要求， 并根据你对技术规范的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    技术规范应包括以下内容：
                    功能定义：概述数据库的核心目标、输入输出及其处理流程。
                    技术要求：列出该数据库的关键技术需求，例如性能指标、可用性、响应时间等。
                    实现方案：提供高层次的实现细节，包括使用的技术框架、算法或工具, 具体到版本。
                    部署要求：说明功能的运行环境需求，如硬件、操作系统、依赖库等。
                    接入方式: 说明用户如何接入该功能，如浏览器接入，微信接入等。

                    在撰写时，请以清晰、结构化的方式逐点展开，确保内容详尽且专业。
                    该功能是产品“{product_name}”的重要组成部分，所以生成的不是技术规范书的整体，而是针对该功能的技术规范段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录。
                    
                    注意：
                    应该以该数据库的技术说明和规范为主, 说明数据库的设计逻辑和核心技术, 多使用英文术语。
                    不要出现段落标题，只写段落的内容即可。
                    以下是产品简述：{product_description}
                    
                    以下是一个简短的示例，这里假设是第4段落，仅供参考，具体子目录要看{index}是多少, 注意，子目录一定要生成:
                    
                    4.1 数据库整体架构
                    分片（Sharding）：将数据水平分割到多个数据库实例中，每个实例负责一部分数据，以提高整体处理能力和存储容量。
                    主从复制（Master-Slave Replication）：设置一个主数据库用于写操作，多个从数据库用于读操作，通过异步或半同步方式复制数据。
                    多主复制（Multi-Master Replication）：允许多个数据库节点同时接受读写操作，适用于高并发写操作的场景。
                    Galera Cluster：基于Galera库的同步多主复制解决方案，提供真正的多主写入和高可用性。
                    4.2 数据库连接与认证
                    连接池（Connection Pooling）：预先创建一组数据库连接，应用程序可以重用这些连接，减少建立和关闭连接的开销。
                    SSL/TLS加密：使用Secure Sockets Layer/Transport Layer Security协议对客户端和服务器之间的通信进行加密。
                    XA事务：支持分布式事务处理，确保跨多个数据库节点的事务的原子性。
                    4.3 数据库性能优化
                    查询优化器（Query Optimizer）：MySQL内置的查询优化器，负责生成最有效的查询执行计划。
                    索引优化：使用B树或哈希索引优化数据检索，特别是针对WHERE子句中使用的列。
                    InnoDB缓冲池（Buffer Pool）：缓存表数据和索引，减少磁盘I/O操作，提高数据访问速度。
                    查询缓存（Query Cache）：缓存SELECT查询的结果，对于相同的查询可以直接返回缓存结果，提高响应速度。
                    4.4 数据安全与备份
                    数据加密（Data Encryption）：使用AES或类似的加密算法对敏感数据进行加密存储。
                    二进制日志（Binary Log）：记录数据库的所有更改，用于数据恢复和复制。
                    物理备份与逻辑备份：物理备份直接复制数据库文件，逻辑备份导出数据内容，两者各有优劣，根据需求选择合适的备份方式。
                    4.5 高可用性与容灾
                    故障转移（Failover）：在主数据库发生故障时，自动将服务切换到备用数据库，确保服务的连续性。
                    数据同步（Data Synchronization）：确保主从数据库之间的数据一致性，使用异步、半同步或全同步方式。
                    数据恢复（Data Recovery）：从备份中恢复数据，包括点时间恢复（Point-in-Time Recovery）和全量恢复。
                    4.6 监控与维护
                    性能监控（Performance Monitoring）：使用慢查询日志、性能模式（Performance Schema）等工具监控数据库性能。
                    自动化运维（Automated Maintenance）：使用工具如pt-online-schema-change进行在线表结构变更，减少维护操作对服务的影响。
                    定期审计（Regular Auditing）：使用审计插件或外部工具定期检查数据库的安全性和性能。
                    """
                    
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index"],
                template=(
                    """
                    {system_message}
                    请描述{product_name}的维护和支持策略，包括维护计划和支持渠道,不少于600字。
                    维护策略应说明{product_name}的维护计划和流程。
                    生成的不是技术规范书的整体, 而是针对该模块的维护与支持段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录
                    以下是产品简述：{product_description}
                    """
                )
            )
            
            
        }

    def generate_text(self, prompt_template, product_description, product_name, index, func_name = None, func_info = None, **kwargs):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == None:
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name
            )
        else:
            # func_name 和 func_info 应该同时不为空，当撰写具体功能技术规范段落时
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name,
                index = index,
                func_name=func_name,
                func_info=func_info
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text
    
    


    def generate_section(self, prompt_template, section_title, index, product_description, product_name, func_name = None, func_info = None, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            product_description=product_description,
            product_name=product_name,
            index = index,
            func_name=func_name,
            func_info=func_info
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string
    
    
    # 提取并保存采购报价单excel
    def save_func_names_json_to_excel(self, func_names_json, output_file_path):
        """
        将 func_names_json 数据处理成 Excel 文件并保存。

        参数:
        func_names_json (list): 包含功能模块和子功能描述的 JSON 数据。
        output_file_path (str): 输出 Excel 文件的路径。
        """
        # Prepare data for DataFrame
        data = []
        for module in func_names_json:
            module_name = module["module_name"]
            for sub_function in module["sub_functions"]:
                sub_function_name = sub_function["sub_function_name"]
                description = sub_function["description"]
                data.append({"功能模块": module_name, "子功能描述": sub_function_name, "具体内容描述": description})

        # Create DataFrame
        df = pd.DataFrame(data)

        # Create a workbook and worksheet
        wb = Workbook()
        ws = wb.active

        # Write DataFrame to worksheet
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=value)

        # Merge cells with the same "功能模块"
        prev_module_name = None
        start_row = 1
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
            current_module_name = row[0].value
            if current_module_name != prev_module_name:
                if prev_module_name is not None:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=row[0].row-1, end_column=1)
                prev_module_name = current_module_name
                start_row = row[0].row
        if prev_module_name is not None:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=ws.max_row, end_column=1)

        # Set alignment
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths with better handling for wide characters
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter  # Get column letter (e.g., 'A', 'B')
            for cell in col:
                if cell.value:
                    # Estimate character width: consider wide characters like Chinese as 2 units
                    value = str(cell.value)
                    adjusted_length = sum(2 if ord(char) > 127 else 1 for char in value)
                    max_length = max(max_length, adjusted_length)
            ws.column_dimensions[col_letter].width = max_length + 2  # Add padding

        # Save the Excel file
        wb.save(output_file_path)

    
    def format_func_info(self, func_info):
        module_name = func_info['module_name']
        sub_functions = func_info['sub_functions']
        
        # 初始化一个空字符串用于存储结果
        result = f"模块名称: {module_name}"
        
        # 遍历子功能并拼装字符串
        for sub_function in sub_functions:
            sub_function_name = sub_function['sub_function_name']
            description = sub_function['description']
            result += f"  子功能名称: {sub_function_name}"
            result += f"    描述: {description}"
        
        return result

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf8') as f:
            product_description = f.read()

        
        # 生成包含功能模块及其子功能描述的JSON
        func_names_json_prompt = f"""
            根据以下数据库服务简述，生成相应的数据库、数据库子功能和子功能描述的JSON。
            以数据库为维度拆分。每个数据库至少包含四个及以上的子功能模块。
            数据库服务简述:
            {product_description}

            生成的JSON应符合以下格式：
            [
                {{
                    "module_name": "数据库1",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }},
                        {{
                            "sub_function_name": "子功能2",
                            "description": "具体内容描述2"
                        }}
                        // 可以添加更多子功能
                    ]
                }},
                {{
                    "module_name": "数据库2",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }}
                        // 可以添加更多子功能
                    ]
                }}
                // 可以添加更多数据库
            ]
            """

        # 使用LLM生成JSON
        func_names_json_response = self.llm.predict(func_names_json_prompt)
        
        try:
            # Strip leading/trailing whitespace and newlines
            stripped_response = func_names_json_response.strip().strip('```json')
            
            # Parse the stripped response into JSON
            func_names_info = json.loads(stripped_response)
            
            # Use func_names_info as needed
            print("数据库详情:", func_names_info)

        except json.JSONDecodeError as e:
            # Handle JSON decoding errors
            print(f"JSON decoding error: {e}")
            func_names_info = None  # Or set to a default value

        except Exception as e:
            # Handle any other unexpected errors
            print(f"An error occurred: {e}")
            func_names_info = None  # Or set to a default value
        try:
            self.save_func_names_json_to_excel(func_names_info, "./采购报价单.xlsx")
            print(f"采购报价单已保存")
        except json.JSONDecodeError as e:
            print(f"JSON 解码错误: {e}")
            func_names_json = None
        
        func_names = [item['module_name'] for item in func_names_info]
        
        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction"]},
        {"title": "数据库服务概述", "prompt": self.prompts["system_overview"]},
        {"title": "数据库系统架构", "prompt": self.prompts["system_architecture"]},
        ]

        # Step 4: 生成前三个部分
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                index = index,
                product_description=product_description,
                product_name=self.product_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

        
        # Step 6: 为每个功能模块生成技术规范
        for index, func_name in enumerate(func_names[:], start=4):  # 从第 4 章开始
            logger.info(f"生成功能技术规范: {func_name}")
            # 查找与 func_name 匹配的模块
            func_info = next((module for module in func_names_info if module['module_name'] == func_name), None)
            
            func_info_str = self.format_func_info(func_info)
            
            if func_info_str is None:
                logger.error(f"未找到功能名称: {func_name}")
                continue
            try:
                section_paragraph = self.generate_section(
                    prompt_template=self.prompts["technical_specification"],
                    section_title=f"{index}. {func_name}技术规范",
                    index = index,
                    product_description=product_description,
                    product_name=self.product_name,
                    func_name=func_name,
                    func_info=func_info_str
                )
                # 清理多余字符并写入文档
                cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
                document.add_paragraph(cleaned_paragraph)
                
            except KeyError as e:
                logger.error(f"KeyError: Missing key '{e.args[0]}' in generate_text method")
                logger.error("Traceback details:", exc_info=True)
            except Exception as e:
                logger.error(f"Error generating technical specification for {product_name}: {e}")
                logger.error("Traceback details:", exc_info=True)

        
        # 添加维护与支持部分
        try:
            maintenance_index = len(func_names) + 4  
            maintenance_paragraph = self.generate_section(
                prompt_template=self.prompts["maintenance_support"],
                section_title=f"{maintenance_index}. 维护与支持",
                index = maintenance_index,
                product_description=product_description,
                product_name=self.product_name,
                func_name=func_name,
                func_info=func_info_str
            )
            cleaned_paragraph = maintenance_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)
        except Exception as e:
            logger.error(f"Error generating maintenance and support section: {e}")
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")
            logger.error("Traceback details:", exc_info=True)
            

if __name__ == "__main__":
    # Define the list of platforms
    platforms = [
        {
            "docx_name": "数据库服务技术规范书.docx",
            "title": "数据库服务技术规范书",
            "product_name": "数据库服务"
        }
        
    ]
    
    # Path to the product overview
    overview_path = "../overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
