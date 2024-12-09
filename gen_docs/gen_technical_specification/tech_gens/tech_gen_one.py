import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os
import json
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
import docx
from typing import Any
import logging
from docx.shared import Pt
import traceback
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment



# 普通技术规范书生成器
# 根据产品简述生成技术规范书



# 从环境变量中获取 OPENAI_API_KEY
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set")
else:
    os.environ["OPENAI_API_KEY"] = str(openai_api_key)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DataPreprocessingModuleSpecGenerator:
    def __init__(self, overview_path, docx_name, title, product_name):
        # 初始化参数
        self.embeddings = OpenAIEmbeddings()
        self.chroma_db = Chroma(embedding_function=self.embeddings)
        self.overview_path = overview_path
        self.docx_name = docx_name
        self.title = title
        # 这里是整个产品的名称
        self.product_name = product_name
        
        # 系统prompt
        self.system_message = """
            你是一个专业的技术文档撰写专家，精通技术规范书的撰写。你的任务是根据输入的产品描述和平台名称，撰写技术规范书。技术规范书的内容必须专业、结构化、详细、准确，并符合以下要求：

            目标明确：清晰表达文档的目标和范围，确保内容针对特定模块的功能、技术细节及使用场景展开。
            条理清晰：文档需组织良好，分章节详细阐述，引言、概述、架构、功能、技术规范等部分明确区分。
            技术性强：提供具体的技术信息，包括系统架构、模块功能、输入输出、处理逻辑和部署要求。
            准确性和可靠性：内容必须严格遵循产品描述中的定义，并符合行业标准或技术最佳实践。
            重要注意事项：
            user_prompt中的产品简述提供该产品的背景知识和功能，帮助你去撰写技术规范书。
            使用专业、正式的语言风格，避免含糊或模糊的描述。
            内容中不要出现*这个符号。
            在描述技术内容时，提供尽可能具体的细节，例如关键组件、数据流和测试标准。
            根据平台类型，结合实际应用场景（例如工业、教育或医疗）撰写内容，体现行业相关性和实用性。
            本次生成是以段落为单位,根据user_prompt去生成对应的段落,而不是整个文档,同时每个段落生成不需要生成生成标题, 例如生成引言段落，不要出现引言标题，直接生成正文即可。
            如果是大模型或相关产品，需要规定GPU或TPU等硬件要求。
            如果是软件功能，请规定数据库设计，接口设计等。
            注意：
            1. 每个段落的目录分级，例如第一个段落引言，需要列出1.1, 1.2等目录，每个段落都需要目录。
            2. 应该使用规定性语言如应该、不得、提供、应支持等，以增强文档的权威性和规范性。
            3. 应该多使用英文术语以提高规范性和权威性。
            4. 如果产品简述中提及是说明性的，那么应该以说明产品的主要功能架构和技术特点为核心。
            
            语言风格示例:
            创建技能组
             提供一个用户友好的界面，允许管理员创建新的技能组。
             管理员可以为每个技能组指定一个唯一的名称和描述，以便于识别和管理。
             技能组的创建过程应支持自定义配置，如服务时间、语言能力、专业领域等。
            管理技能组
             管理员可以编辑已存在的技能组信息，包括修改名称、描述和配置参数。
             提供技能组的启用和禁用功能，允许临时停用某个技能组而不删除其配置。
             技能组配置应支持版本控制，记录历史更改并允许恢复到之前的配置。

            
            """
        
        # 各个部分的PromptTemplate
        self.prompts = {
            "introduction": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请编写一段关于{product_name}的技术规范书引言，简要介绍文档的目的和范围，不少于1500字。
                    请包括文档的目标、涵盖的范围以及相关的参考文档和定义。
                    生成的不是技术规范书的整体, 而是针对该模块的引言段落。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    1.1 目的
                    本文档的编制目的是为了定义智能客服系统的技术标准和规范，确保系统开发、部署及运维过程的规范性和一致性。通过对系统的功能、性能、接口、设计和实施要求的详细规定，旨在保障系统能够满足业务需求，同时提升系统的稳定性、安全性和用户体验。本文档将作为项目开发团队、测试团队、运维团队以及相关利益相关者的参考依据，确保各方对系统的理解和期望达成一致。
                    1.2 范围
                    本技术规范书全面覆盖智能客服系统的构建和运行所需的全部技术要求。具体包括但不限于：
                     系统架构设计，包括系统的总体设计思路、模块划分和数据流向。
                     功能实现细节，涉及系统的核心功能点以及附加功能的具体实现方式。
                     接口标准规定，明确系统与外部系统集成时的接口规范和数据交互格式。
                     安全性要求，确保系统在数据存储、传输和处理过程中的安全性。
                     可靠性指标，保障系统的高可用性、故障恢复能力和性能稳定性。
                     维护和支持指南，提供系统的维护流程、升级策略和技术支持渠道。
                    1.3 参考文档
                    为确保本技术规范书的科学性和权威性，编制过程中参考了以下文档和标准：
                     ISO/IEC 25010:2011 系统和软件工程——系统和软件产品质量模型，提供了系统和软件质量的通用评价框架。
                     相关API接口文档，为系统提供了与外部服务和系统集成所需的接口信息和数据交互标准。
                     数据保护法规，包括GDPR、CCPA等，确保系统在处理用户数据时遵守相应的法律法规。
                    1.4 定义和缩略语
                    为了确保本文档中的术语使用一致和清晰，以下列出了一些关键词汇及其定义：
                     CRM（Customer Relationship Management）：客户关系管理系统，用于帮助企业管理与客户的互动信息。
                     API（Application Programming Interface）：应用程序编程接口，允许不同软件组件之间进行互相通信。
                     SDK（Software Development Kit）：软件开发工具包，提供了一套工具、指南和API，以便开发者创建软件应用程序。
                     UI（User Interface）：用户界面，指系统与用户交互的界面部分，包括布局、视觉元素和交互逻辑。
                     UX（User Experience）：用户体验，指用户在使用系统过程中的总体感受和体验。
                     HTTPS（Hypertext Transfer Protocol Secure）：安全超文本传输协议，通过对传输数据进行加密，提供了数据传输的安全性保障。
                     SSL/TLS（Secure Sockets Layer / Transport Layer Security）：安全套接层/传输层安全，两种用于在互联网上确保数据传输安全的协议。
                     NLP（Natural Language Processing）：自然语言处理，指使计算机能够理解和处理人类语言的技术。

                    """
                )
            ),
            "system_overview": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请提供一个{product_name}的系统概述，包括其主要功能和用途, 不少于2000字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统概述段落，是整个技术规范书的第二段，第一段是引言，所以专注于系统概述本身，不需要涉及引言或系统架构等段落。
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    2.1 系统功能
                    模型剪枝系统旨在提供高效、灵活的模型优化解决方案，通过整合多种剪枝技术和工具，简化模型压缩流程。以下是系统的主要功能：

                    2.1.1 多模型支持
                    深度学习框架：支持TensorFlow、PyTorch等多种主流深度学习框架。
                    模型格式：兼容ONNX、HDF5等多种模型文件格式。
                    预训练模型：提供预训练模型的剪枝支持，加速用户操作。
                    2.1.2 剪枝技术
                    结构化剪枝：支持通道、层等结构化剪枝方法。
                    非结构化剪枝：支持权重级别的非结构化剪枝。
                    自动化剪枝：根据用户设定的阈值和策略，自动执行剪枝操作。
                    2.1.3 剪枝优化
                    性能评估：提供剪枝前后的模型性能对比分析。
                    优化建议：根据模型特性和剪枝结果，提供优化建议。
                    稀疏化处理：对剪枝后的模型进行稀疏化处理，提高运行效率。
                    2.1.4 用户交互
                    图形界面：提供直观易用的图形用户界面，简化操作流程。
                    命令行工具：支持命令行操作，方便批量处理和自动化脚本编写。
                    日志记录：详细记录剪枝过程中的所有操作和结果，便于追踪和调试。
                    2.2 用户特征
                    模型剪枝系统面向以下用户群体：

                    研究人员：使用系统进行模型优化实验，探索新的剪枝技术和方法。
                    开发人员：部署和维护模型剪枝流程，优化模型性能。
                    数据科学家：利用系统提高模型训练效率和预测准确性。
                    2.3 运行环境
                    模型剪枝系统的运行环境要求如下：

                    服务器端
                    操作系统：支持Linux和Windows服务器操作系统。
                    部署：支持物理服务器、虚拟化环境及云服务平台部署。
                    依赖库：兼容主流的深度学习库和工具，如TensorFlow、PyTorch等。
                    客户端
                    操作系统：支持Windows、macOS和Linux桌面操作系统。
                    浏览器：确保与主流浏览器如Chrome、Firefox、Safari和Edge兼容。
                    开发工具：支持集成开发环境（IDE）和Jupyter Notebook等交互式编程工具。


                    """
                )
            ),
            "system_architecture": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name"],
                template=(
                    """
                    {system_message}
                    请详细描述{product_name}的系统架构, 不少于2500字。
                    生成的不是技术规范书的整体, 而是针对该产品的系统架构段落。
                    架构描述应包括数据流、处理步骤和关键组件。
                    根据你对系统架构的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    以下是产品简述：{product_description}
                    
                    以下是示例:
                    3.1 总体架构
                    模型剪枝系统将采用一种模块化、分布式的架构设计，以确保系统的可扩展性、可维护性和高可用性。系统的主要架构分为以下几个层次：

                    3.1.1 剪枝策略层
                    剪枝策略层负责定义和实施模型剪枝的各种策略。该层主要包括：

                    剪枝算法选择：为不同的模型和任务选择合适的剪枝算法。
                    剪枝比例配置：根据模型性能和需求，配置剪枝的比例和阈值。
                    剪枝规则定义：定义剪枝的具体规则，如基于权重大小、梯度信息等。
                    3.1.2 剪枝执行层
                    剪枝执行层是系统的核心，负责实际执行模型剪枝操作。包括但不限于：

                    模型分析：分析模型的结构和参数，确定可剪枝的部分。
                    剪枝操作：根据剪枝策略，对模型进行实际的剪枝操作。
                    模型验证：验证剪枝后的模型性能，确保满足要求。
                    3.1.3 结果输出层
                    结果输出层负责输出剪枝后的模型和相关报告。该层主要包括：

                    剪枝后模型存储：存储剪枝后的模型文件，供后续使用。
                    剪枝报告生成：生成剪枝过程的详细报告，包括剪枝前后的模型对比、性能变化等。
                    3.2 系统模块划分
                    系统根据功能需求被划分为多个模块，以实现职责的分离和功能的独立。

                    3.2.1 模型输入模块
                    模型加载：负责加载原始模型文件。
                    数据准备：准备用于剪枝分析和验证的数据集。
                    3.2.2 剪枝处理模块
                    剪枝算法实现：实现具体的剪枝算法，如权重剪枝、通道剪枝等。
                    剪枝优化：对剪枝后的模型进行优化，如稀疏化处理、量化等。
                    性能评估：评估剪枝后模型的性能，如精度、速度等。
                    3.2.3 结果输出模块
                    模型导出：导出剪枝后的模型文件，支持多种格式。
                    报告生成：生成详细的剪枝报告，包括剪枝详情、性能对比等。
                    
                    """
                )
            ),
            
            "technical_specification": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index", "func_name", "func_info"],
                template=(
                    """
                    {system_message}
                    请详细描述功能“{func_name}”的技术规范段落, 不少于3000字。
                    这是该功能的具体子模块和描述，请参考：“{func_info}”
                    技术规范需要涵盖该功能的所有子模块技术要求， 并根据你对技术规范的理解，请尽量详细描述，进行技术和架构上的扩展。
                    
                    技术规范应包括以下内容：
                    功能定义：概述功能的核心目标、输入输出及其处理流程。
                    技术要求：列出该功能的关键技术需求，例如性能指标、可用性、响应时间等。
                    实现方案：提供高层次的实现细节，包括使用的技术框架、算法或工具, 具体到版本。
                    数据库设计和接口设计。如果是平台类或功能类产品, 应根据架构设计,详细描述数据库表设计、接口设计、API设计等， 
                    注意，如果不是后端功能，则不要涉及数据库设计等内容。
                    如果是大模型相关技术，请详细说明和规范模型相关的核心技术和设计逻辑。
                    部署要求：说明功能的运行环境需求，如硬件、操作系统、依赖库等。
                    接入方式: 说明用户如何接入该功能，如浏览器接入，微信接入等。

                    在撰写时，请以清晰、结构化的方式逐点展开，确保内容详尽且专业。
                    该功能是产品“{product_name}”的重要组成部分，所以生成的不是技术规范书的整体，而是针对该功能的技术规范段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录。
                    
                    不要出现段落标题，只写段落的内容即可。
                    以下是产品简述：{product_description}
                    
                    以下是两个示例，请根据产品的类别去参考。
                    以下是一个平台类的示例，这里假设是第4段落，仅供参考，具体子目录要看{index}是多少, 注意，子目录一定要生成:
                    4.1 桌面网站接入
                    4.1.1 JavaScript SDK
                     提供JavaScript SDK以供集成到企业桌面网站中，该SDK应包含以下功能：
                     用户身份验证：确保用户与系统之间的通信是安全的。
                     会话管理：允许创建、维护和结束用户会话。
                     消息发送和接收：支持文本、图片、文件等多种消息类型的发送和接收。
                     事件监听：能够响应用户交互事件，如点击、输入等。
                     UI组件：提供标准的聊天窗口UI组件，支持自定义样式以符合企业品牌形象。
                    4.1.2 通讯机制
                     系统应支持WebSocket协议，以实现实时、双向的通信。
                     对于不支持WebSocket的环境，系统应提供长轮询的备选方案，以保证消息的可靠传输。
                     所有通讯应通过HTTPS进行加密，以确保数据传输的安全性。
                    4.2 移动网站接入
                    4.2.1 自适应设计
                     移动端接入应使用响应式设计，确保UI组件能够根据不同屏幕尺寸进行适配。
                     SDK应提供与桌面网站接入相同的功能，并优化触摸操作和网络条件变化的响应。
                    4.3 数据库表设计
                     用户表设计：
                        CREATE TABLE users (
                            `user_id` INT AUTO_INCREMENT PRIMARY KEY,
                            `username` VARCHAR(255) UNIQUE NOT NULL,
                            `password_hash` VARCHAR(255) NOT NULL,
                            `role` ENUM('admin', 'agent', 'customer') NOT NULL
                        );

                    以下是一个模型类的示例，这里假设是第4段落，仅供参考，具体子目录要看{index}是多少, 注意，子目录一定要生成:
                    4.1 算法集成与优化
                    4.1.1 时间序列分析算法包

                    提供一套综合的时间序列分析算法包，专为集成至企业级应用设计。
                    包含多种先进的预测模型，如ARIMA、LSTM、Prophet等，并支持模型间的无缝切换。
                    提供模型参数自动调优功能，利用网格搜索、贝叶斯优化等技术寻找最优参数组合。
                    实现模型的增量学习机制，以便在不丢失历史学习成果的基础上适应新数据。
                    4.1.2 实时预测与反馈循环

                    设计高效的实时预测模块，能够迅速处理流式数据并给出即时预测结果。
                    构建闭环反馈系统，允许模型根据实时预测误差自我调整和学习。
                    4.2 性能优化策略
                    4.2.1 算法并行化处理

                    利用多核CPU和GPU加速计算密集型任务，提升算法运行效率。
                    探索分布式计算框架的应用，以支持大规模时间序列数据的处理。
                    4.2.2 内存管理与资源优化

                    优化算法内存使用，减少不必要的数据存储和复制操作。
                    设计智能缓存机制，提高数据访问速度和处理吞吐量。
                    4.3 模型评估与监控
                    4.3.1 先进的评估指标

                    采用多种评估指标全面衡量模型性能，包括RMSE、MAE、MAPE等。
                    引入回测机制，在历史数据上模拟预测场景，验证模型的稳健性和可靠性。
                    4.3.2 实时监控与告警系统

                    构建实时监控平台，跟踪模型在实际运行中的表现。
                    设置合理的告警阈值，一旦发现性能下降或异常立即触发响应措施。
                    4.4 可扩展性与兼容性
                    4.4.1 模块化设计

                    遵循模块化设计原则，方便后期功能扩展和维护升级。
                    提供清晰的API接口文档，促进与其他系统的集成与合作。
                    4.4.2 跨平台适配

                    确保算法包能在多种操作系统及硬件环境下稳定运行。
                    针对不同编程语言提供友好的绑定和支持。
                    """
                    
                )
            ),
            "maintenance_support": PromptTemplate(
                input_variables=["system_message", "product_description", "product_name", "index"],
                template=(
                    """
                    {system_message}
                    请描述{product_name}的维护和支持策略，包括维护计划和支持渠道,不少于400字。
                    维护策略应说明{product_name}的维护计划和流程。
                    生成的不是技术规范书的整体, 而是针对该模块的维护与支持段落。
                    这是整个技术规范书的第{index}段， 子目录应该是{index}.1, {index}.2, {index}.3依次类推，请确保段落的目录分级, 要划分子目录
                    以下是产品简述：{product_description}
                    """
                )
            )
            
            
        }

    def generate_text(self, prompt_template, product_description, product_name, index, func_name = None, func_info = None, **kwargs):
        """
        Generate text using LLMChain with the provided prompt template.
        """
        # Define the system message
        system_message = SystemMessage(content=self.system_message)

        # Create the user message using the PromptTemplate
        if func_name == None:
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name
            )
        else:
            # func_name 和 func_info 应该同时不为空，当撰写具体功能技术规范段落时
            
            user_prompt = prompt_template.format(
                system_message=self.system_message,
                product_description=product_description,
                product_name=product_name,
                index = index,
                func_name=func_name,
                func_info=func_info
            )
        user_message = HumanMessage(content=user_prompt)

        # Send both system and user messages to the LLM
        response = self.llm.predict_messages([system_message, user_message])

        generated_text = response.content
        logger.info("Generated Text: %s", generated_text)
        return generated_text
    
    


    def generate_section(self, prompt_template, section_title, index, product_description, product_name, func_name = None, func_info = None, **kwargs):
        """
            Generate the specified section of text and add a title and number.
        """
        
        section_content = self.generate_text(
            prompt_template=prompt_template,
            product_description=product_description,
            product_name=product_name,
            index = index,
            func_name=func_name,
            func_info=func_info
        )
        
        # Create a new paragraph with the title and content
        section_string = f"{section_title}\n{section_content}"
    
        return section_string
    
    
    # 提取并保存采购报价单excel
    def save_func_names_json_to_excel(self, func_names_json, output_file_path):
        """
        将 func_names_json 数据处理成 Excel 文件并保存。

        参数:
        func_names_json (list): 包含功能模块和子功能描述的 JSON 数据。
        output_file_path (str): 输出 Excel 文件的路径。
        """
        # Prepare data for DataFrame
        data = []
        for module in func_names_json:
            module_name = module["module_name"]
            for sub_function in module["sub_functions"]:
                sub_function_name = sub_function["sub_function_name"]
                description = sub_function["description"]
                data.append({"功能模块": module_name, "子功能描述": sub_function_name, "具体内容描述": description})

        # Create DataFrame
        df = pd.DataFrame(data)

        # Create a workbook and worksheet
        wb = Workbook()
        ws = wb.active

        # Write DataFrame to worksheet
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=value)

        # Merge cells with the same "功能模块"
        prev_module_name = None
        start_row = 1
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
            current_module_name = row[0].value
            if current_module_name != prev_module_name:
                if prev_module_name is not None:
                    ws.merge_cells(start_row=start_row, start_column=1, end_row=row[0].row-1, end_column=1)
                prev_module_name = current_module_name
                start_row = row[0].row
        if prev_module_name is not None:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=ws.max_row, end_column=1)

        # Set alignment
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths with better handling for wide characters
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter  # Get column letter (e.g., 'A', 'B')
            for cell in col:
                if cell.value:
                    # Estimate character width: consider wide characters like Chinese as 2 units
                    value = str(cell.value)
                    adjusted_length = sum(2 if ord(char) > 127 else 1 for char in value)
                    max_length = max(max_length, adjusted_length)
            ws.column_dimensions[col_letter].width = max_length + 2  # Add padding

        # Save the Excel file
        wb.save(output_file_path)

    
    def format_func_info(self, func_info):
        module_name = func_info['module_name']
        sub_functions = func_info['sub_functions']
        
        # 初始化一个空字符串用于存储结果
        result = f"模块名称: {module_name}"
        
        # 遍历子功能并拼装字符串
        for sub_function in sub_functions:
            sub_function_name = sub_function['sub_function_name']
            description = sub_function['description']
            result += f"  子功能名称: {sub_function_name}"
            result += f"    描述: {description}"
        
        return result

    def process(self):
        self.llm = ChatOpenAI(model="gpt-4o-mini")

        # 读取产品简述
        with open(self.overview_path, 'r', encoding='utf8') as f:
            product_description = f.read()

        
        # 生成包含功能模块及其子功能描述的JSON
        func_names_json_prompt = f"""
            根据以下产品简述，生成相应的功能模块、子模块和子功能描述的JSON。
            一个产品应有多个功能模块，每个功能模块包含多个子模块，一个子模块对应一个子模块描述。
            至少生成5个及以上功能模块, 每个功能模块至少包含6个及以上的子功能模块。
            产品简述:
            {product_description}

            生成的JSON应符合以下格式：
            [
                {{
                    "module_name": "模块1",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }},
                        {{
                            "sub_function_name": "子功能2",
                            "description": "具体内容描述2"
                        }}
                        // 可以添加更多子功能
                    ]
                }},
                {{
                    "module_name": "模块2",
                    "sub_functions": [
                        {{
                            "sub_function_name": "子功能1",
                            "description": "具体内容描述1"
                        }}
                        // 可以添加更多子功能
                    ]
                }}
                // 可以添加更多模块
            ]
            """

        # 使用LLM生成JSON
        func_names_json_response = self.llm.predict(func_names_json_prompt)
        
        try:
            # Strip leading/trailing whitespace and newlines
            stripped_response = func_names_json_response.strip().strip('```json')
            
            # Parse the stripped response into JSON
            func_names_info = json.loads(stripped_response)
            
            # Use func_names_info as needed
            print("功能详情:", func_names_info)

        except json.JSONDecodeError as e:
            # Handle JSON decoding errors
            print(f"JSON decoding error: {e}")
            func_names_info = None  # Or set to a default value

        except Exception as e:
            # Handle any other unexpected errors
            print(f"An error occurred: {e}")
            func_names_info = None  # Or set to a default value
        try:
            self.save_func_names_json_to_excel(func_names_info, "./采购报价单.xlsx")
            print(f"采购报价单已保存")
        except json.JSONDecodeError as e:
            print(f"JSON 解码错误: {e}")
            func_names_json = None
        
        func_names = [item['module_name'] for item in func_names_info]
        
        # 创建Word文档并添加标题
        document = docx.Document()
        document.add_heading(self.title, level=1)

        sections = [
        {"title": "引言", "prompt": self.prompts["introduction"]},
        {"title": "系统概述", "prompt": self.prompts["system_overview"]},
        {"title": "系统架构", "prompt": self.prompts["system_architecture"]},
        ]

        # Step 4: 生成前三个部分
        for index, section in enumerate(sections, start=1):
            section_paragraph = self.generate_section(
                prompt_template=section["prompt"],
                section_title=f"{index}. {section['title']}",
                index = index,
                product_description=product_description,
                product_name=self.product_name
            )
            # 清理多余字符并写入文档
            cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)

        
        # Step 6: 为每个功能模块生成技术规范
        for index, func_name in enumerate(func_names[:], start=4):  # 从第 4 章开始
            logger.info(f"生成功能技术规范: {func_name}")
            # 查找与 func_name 匹配的模块
            func_info = next((module for module in func_names_info if module['module_name'] == func_name), None)
            
            func_info_str = self.format_func_info(func_info)
            
            if func_info_str is None:
                logger.error(f"未找到功能名称: {func_name}")
                continue
            try:
                section_paragraph = self.generate_section(
                    prompt_template=self.prompts["technical_specification"],
                    section_title=f"{index}. {func_name}技术规范",
                    index = index,
                    product_description=product_description,
                    product_name=self.product_name,
                    func_name=func_name,
                    func_info=func_info_str
                )
                # 清理多余字符并写入文档
                cleaned_paragraph = section_paragraph.replace('', '').replace('#', '')
                document.add_paragraph(cleaned_paragraph)
                
            except KeyError as e:
                logger.error(f"KeyError: Missing key '{e.args[0]}' in generate_text method")
                logger.error("Traceback details:", exc_info=True)
            except Exception as e:
                logger.error(f"Error generating technical specification for {product_name}: {e}")
                logger.error("Traceback details:", exc_info=True)

        
        # 添加维护与支持部分
        try:
            maintenance_index = len(func_names) + 4  
            maintenance_paragraph = self.generate_section(
                prompt_template=self.prompts["maintenance_support"],
                section_title=f"{maintenance_index}. 维护与支持",
                index = maintenance_index,
                product_description=product_description,
                product_name=self.product_name,
                func_name=func_name,
                func_info=func_info_str
            )
            cleaned_paragraph = maintenance_paragraph.replace('', '').replace('#', '')
            document.add_paragraph(cleaned_paragraph)
        except Exception as e:
            logger.error(f"Error generating maintenance and support section: {e}")
            
            
        # 保存文档
        document.save(self.docx_name)
        logger.info(f"{self.product_name}技术规范书已保存为 {self.docx_name}")



def generate_all_platform_docs(overview_path, platforms):
    """
    Generates technical specification documents for multiple platforms.

    Args:
        overview_path (str): Path to the overview text file.
        platforms (list of dict): List of platform configurations, where each dictionary
                                  contains 'docx_name', 'title', and 'product_name'.
    """
    for platform in platforms:
        logger.info(f"Starting generation for {platform['product_name']}...")
        generator = DataPreprocessingModuleSpecGenerator(
            overview_path=overview_path,
            docx_name=platform["docx_name"],
            title=platform["title"],
            product_name=platform["product_name"]
        )
        try:
            generator.process()
            logger.info(f"Successfully generated document for {platform['product_name']}!")
        except Exception as e:
            logger.error(f"Failed to generate document for {platform['product_name']}: {e}")
            logger.error("Traceback details:", exc_info=True)
            

if __name__ == "__main__":
    # Define the list of platforms
    platforms = [
        {
            "docx_name": "虚拟仿真技术支持的智能产品研发优化管理平台技术规范书.docx",
            "title": "虚拟仿真技术支持的智能产品研发优化管理平台技术规范书",
            "product_name": "虚拟仿真技术支持的智能产品研发优化管理平台"
        }
        
    ]
    
    # Path to the product overview
    overview_path = "./overview.txt"
    
    # Generate documents for all platforms
    generate_all_platform_docs(overview_path, platforms)
